# -*- coding: utf-8 -*-
"""Build the SMPL: VI Treatment Vessel Setup Unit Test Script (.docx) from the AZ Test
Script Template, driven by the XStep functional spec + mock-up. Fills the template's
metadata tables (Referenced Documents, Test Script header, Pre-requisites / Set-Up) and
replaces the step table with Treatment-Vessel-specific Action / Expected-Result steps.

Reference docs (in the UT Scripts folder): 'AZ Test Script Template (2).docx' (format) and
'Phase 1 Example - SMPL_ Component Goods Issue (1).docx' (step style / granularity)."""
import sys, io, copy, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor

UTDIR=r'../AZ Phase 3 Virus Inactivation and filtration/UT Scripts'
TEMPLATE=UTDIR+'/AZ Test Script Template (2).docx'
OUT=UTDIR+'/SMPL_VI Treatment Vessel Setup Unit Test Script.docx'
BLACK=RGBColor(0,0,0)

doc=Document(TEMPLATE)

def _mkrun(p, text):
    """Add a run and force BLACK, overriding the template's blue instructional rPr."""
    r=p.add_run(text); r.font.color.rgb=BLACK
    # also strip any inherited color on the paragraph mark so empty runs don't show blue
    return r

def setc(cell, text):
    """Replace a cell's content with text (multi-line on \\n), preserving the paragraph style, in black."""
    for p in cell.paragraphs[1:]:
        p._p.getparent().remove(p._p)
    p=cell.paragraphs[0]
    for r in list(p.runs): r._r.getparent().remove(r._r)
    lines=str(text).split('\n')
    _mkrun(p, lines[0])
    for ln in lines[1:]:
        np=cell.add_paragraph(); _mkrun(np, ln)

def replace_para(old, new):
    for p in doc.paragraphs:
        if p.text.strip()==old:
            for r in list(p.runs): r._r.getparent().remove(r._r)
            _mkrun(p, new); return True
    return False

# ---- title page ----
replace_para('<Project Name>', 'AZ Phase 3 - Virus Inactivation (AZD0543)')

# ---- Table 0: Document Approval / Review (Author row; Reviewed/Approved signed at review) ----
t0=doc.tables[0]
setc(t0.rows[2].cells[1], 'Carlos Redekopp / Author')
setc(t0.rows[2].cells[3], '23JUL2026')

# ---- Table 1: Version Control ----
t1=doc.tables[1]
setc(t1.rows[2].cells[0], '1.0')
setc(t1.rows[2].cells[1], '23JUL2026')
setc(t1.rows[2].cells[2], 'Initial Document')
setc(t1.rows[2].cells[3], 'Carlos Redekopp')

# ---- Table 2: Referenced Documents (No. / Title & ID / Version / Location) ----
t2=doc.tables[2]
refs=[('1','SMPL: VI Treatment Vessel Setup XStep Design Specification','1.0',
       'AZ Phase 3 Virus Inactivation and filtration / Functional Specs'),
      ('2','AZD0543 - Virus Inactivation (Low pH) Manufacturing Directions (PN 8012441)','',''),
      ('3','SiMPL XStep Library - Process Manufacturing (DE1 100)','',''),
      ('4','AZ Phase 3 VI + CDF EBR (PI Sheet) mock-up','','AZ Phase 3 Virus Inactivation and filtration')]
# t2 rows: 0 header, 1 col-headers, 2 & 3 blank data, 4 "Add more rows"
addmore=t2.rows[4]._tr
tmpl_tr=copy.deepcopy(t2.rows[2]._tr)
# fill the two existing blank rows, insert clones before the "Add more rows" row for the rest
data_rows=[t2.rows[2], t2.rows[3]]
for _ in range(len(refs)-2):
    nt=copy.deepcopy(tmpl_tr); addmore.addprevious(nt)
from docx.table import _Row
data_rows=[r for r in t2.rows[2:] if r._tr is not addmore][:len(refs)]
for row,(no,title,ver,loc) in zip(data_rows, refs):
    setc(row.cells[0],no); setc(row.cells[1],title); setc(row.cells[2],ver); setc(row.cells[3],loc)

# ---- Table 4: Test Script header ----
t4=doc.tables[4]
setc(t4.rows[0].cells[1], 'SMPL: VI Treatment Vessel Setup V1.0') # Test Script ID & Version #
setc(t4.rows[0].cells[3], 'Unit Test')                           # Test Phase
setc(t4.rows[1].cells[1], 'SMPL: VI Treatment Vessel Setup')     # Test Name
setc(t4.rows[2].cells[1],                                        # Test Objective
     'To verify that the SMPL: VI Treatment Vessel Setup XStep correctly captures the VI Treatment Vessel set-up: '
     'driving conditional activation of the Bag / Tank / Vent Filter sections from the Yes/No gate dropdowns, '
     'validating the material, batch, autoclave-date and equipment entries, posting the Z_PICONS goods issue for the '
     'consumed bag and filters, and enforcing the per-row Performed By and step-level Recorded By electronic signatures.')
setc(t4.rows[3].cells[1],                                        # Requirements Tested
     'The test confirms that the gate dropdowns activate/deactivate the correct sections via the TABLE / PROC_INSTR '
     'ACTIVATE / DEACTIVATE commands (LV_DROP), that the Part No. -> Material Description and Batch No. -> Exp. Date '
     'input validations populate correctly, that the autoclave-expiry validation, equipment retrieval and pressure '
     'date/time stamp behave as designed, that the Z_PICONS process message (movement type 261) is posted on signing, '
     'and that the Performed By and Recorded By signatures enforce traceable completion.')

# ---- Table 5: Pre-requisites / Set-Up / Data ----
t5=doc.tables[5]
setc(t5.rows[1].cells[0],
     'Valid Process Order: A process order must exist in SAP with the SMPL: VI Treatment Vessel Setup XStep correctly '
     'assigned via the relevant master recipe.\n'
     'System Configuration: The process message Z_PICONS communicates goods issue (movement type 261) to SAP; the '
     'CT04 characteristic ZSMPL_CHAR_YES_NO (Yes / No) exists and backs the gate, attachment and filters-attached '
     'dropdowns.\n'
     'Master Data: Test materials (bag, product filter, vent filter) exist in the material master with valid batches '
     '(shelf-life expiry) in the batch master; a scale / balance is maintained as assigned equipment (EQTYP B) with a '
     'calibration due date.')
setc(t5.rows[3].cells[0],
     'Input Parameter Setup: In SiMPL MBR, populate the input parameters with test values - IV_HEADER / IV_INSTR left '
     'as DEFAULT (or authored text); IV_ACTIVE / IV_ACT_FLP set as required to exercise the activation logic; the gate '
     'dropdowns (Use Affinity Vessel? / Bag Required? / Tank Required? / Vent Filter Required?) available for selection.')
setc(t5.rows[3].cells[1],
     'A test bag (part / batch), a product filter (part / batch), a vent filter (part / batch, using a bolded '
     'autoclave-required part to exercise the autoclave fields), a batch whose expiry is in the past (negative test), '
     'and an assigned scale / balance equipment ID with a calibration due date.')

# ---- Table 6: System / Environment (tester name / date / run # completed at execution) ----
t6=doc.tables[6]
setc(t6.rows[0].cells[1], 'DE1 V1')   # System Name & Version #
setc(t6.rows[1].cells[1], 'DE1')      # Environment

# ---- Table 7: Test Steps ----
STEPS=[
 ('Create a Process Order that contains the SMPL: VI Treatment Vessel Setup XStep.',
  'Process Order created containing SMPL: VI Treatment Vessel Setup.'),
 ('Open the EBR for the selected Process Order / Control Recipe and navigate to SMPL: VI Treatment Vessel Setup.',
  'The EBR opens successfully without any errors and the SMPL: VI Treatment Vessel Setup XStep is visible.'),
 ('Verify the values entered in parameter IV_HEADER and parameter IV_INSTR in SiMPL MBR appear in the XStep header / instruction.',
  'Header and instruction text are displayed as authored in SiMPL MBR.'),
 ('Select "Use Affinity Product Vessel as VI Treatment Vessel?" = Yes.',
  'The whole step is deactivated (marked N/A) - the step’s DEACTIVATE commands fire on the gate value (LV_DROP).'),
 ('Select "Use Affinity Product Vessel as VI Treatment Vessel?" = No.',
  'The step remains active and the vessel set-up sections are available for entry.'),
 ('Select "Bag Required?" = No, then = Yes.',
  'No (LV_DROP = 2) fires TABLE.DEACTIVATE so the Bag Information table is deactivated; Yes (LV_DROP <> 2) fires '
  'TABLE.ACTIVATE so it activates for entry.'),
 ('In the Bag Information table, enter the Bag Part No.',
  'The Material Description is populated automatically and displayed read-only (ZSMPL_FM_GET_MAT_DETAIL).'),
 ('Enter the Bag Batch No. and the Serial No.',
  'The Exp. Date is populated automatically from the batch (ZSMPL_FM_GET_EXPIRY_DATE); the Serial No. is accepted.'),
 ('Select Yes for the Mixer Drive attached?, Top Base attached? and Magnetic Clamp attached? dropdowns.',
  'The dropdown selections are accepted (ZSMPL_CHAR_YES_NO).'),
 ('Enter a username in the "Performed By" field on the bag line, then enter the password in the pop-up window.',
  'The signature is accepted and the bag line is completed.'),
 ('Select "Tank Required?" = Yes and enter the Tank ID, CIP Batch ID, Pressure Test ID and SIP Batch ID.',
  'The Tank Information table activates; the entries are accepted and the CIP Exp. Date/Time is populated.'),
 ('Press "Get Pressure Date/Time" in the Tank Information section.',
  'The Pressure Test Date and Pressure Test Time are stamped and displayed read-only (/SMPL/PPPI_FM_GET_DATE_TIME).'),
 ('Select "Vent Filter Required?" = Yes; in the Vent Filter table enter the Part No., Batch No., Serial No. and Autoclave Cycle No.',
  'The Vent Filter table activates; the Material Description and Exp. Date populate; the entries are accepted.'),
 ('Enter an Autoclave Exp. Date with an invalid month (e.g. 2027.13.01).',
  'The entry is rejected with "Incorrect date format entered - use yyyy.mm.dd" (ZSMPL_FM_CHECK_CHAR_DATE).'),
 ('Enter "N/A" for the Autoclave Exp. Date (where no autoclave expiry applies).',
  'The entry is accepted.'),
 ('Enter the "Performed By" signature on the vent filter line.',
  'The signature is accepted and the vent filter line is completed.'),
 ('In the Product Filter table, enter a product filter line (Part No., Batch No., Serial No., Autoclave Cycle No., '
  'Autoclave Exp.) and press Add Row to add a second product filter.',
  'A new row is added and numbered automatically (ZSMPL_FM_CUSTOM_INDEX); both lines accept entry.'),
 ('Enter a Batch No. whose expiration date is in the past on a bag / filter line.',
  'The entry is rejected with EXPIRY_DATE_IN_PAST (ZSMPL_FM_GET_EXPIRY_DATE).'),
 ('Press "Get Scale / Balance".',
  'The assigned balance ID and its Calibration Due Date are retrieved and displayed (ZSMPL_FM_GET_ASSIGNED_EQUI_EBR).'),
 ('Enter the Tare Weight of the VI Treatment Vessel.',
  'The value is accepted.'),
 ('Select the "Product Filter attached?", "Vent Filter attached?" and "No filters - vessel is a bag?" dropdowns.',
  'The selections are accepted.'),
 ('Attempt to complete the step without entering the Recorded By / Date signature.',
  'The step cannot be completed - the Recorded By signature is mandatory.'),
 ('Enter the Recorded By / Date signature (username in the field, password in the pop-up window).',
  'The signature is accepted and the step can be completed.'),
 ('Navigate to transaction CO54 and verify the Z_PICONS process messages for the consumed bag and filters on the current order.',
  'Process messages exist with the correct material, batch, quantity and movement type (261).'),
 ('Re-open the completed XStep.',
  'The recorded values, populated outputs, the pressure date/time stamp and the signatures are displayed and retained.'),
]
t7=doc.tables[7]
row_tmpl=copy.deepcopy(t7.rows[1]._tr)        # a bordered data row to clone
for r in list(t7.rows[1:]):                    # remove the 6 template data rows
    r._tr.getparent().remove(r._tr)
for i,(act,exp) in enumerate(STEPS, start=1):
    nt=copy.deepcopy(row_tmpl); t7._tbl.append(nt)
    row=t7.rows[-1]
    setc(row.cells[0], str(i)); setc(row.cells[1], act); setc(row.cells[2], exp)
    setc(row.cells[3], ''); setc(row.cells[4], '')

# ---- global: force every coloured run to black (kill the template's blue instructional text) ----
for part in (doc.element.body,):
    for col in part.iter(qn('w:color')):
        col.set(qn('w:val'), '000000')
        for a in ('themeColor', 'themeTint', 'themeShade'):
            if col.get(qn('w:'+a)) is not None:
                del col.attrib[qn('w:'+a)]

doc.save(OUT)
sys.stdout.flush()
print('built', OUT, '|', len(STEPS), 'steps')
