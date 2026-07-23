# -*- coding: utf-8 -*-
"""Build the SMPL: WFI Container Setup & Charge XStep Design Specification (.docx) in
the AZ template. From-scratch build (AZ Phase 3 mock-up only): clones the Three Variable
Calc doc for the shell (title page, styles, live TOC), rebuilds the 13 sections from the
WFI Container Setup & Charge mock-up, and reuses the Phase-2 FM/pseudocode emitters.

The step sets up the Day 1 / Day 2 WFI collection vessel: a Data Input table records the
bag and WFI filter (Item dropdown, Part No., Batch No.->Exp. Date input validation,
Performed By per line, Z_PICONS goods issue on signing), then footer fields record the
WFI valve ID, stamp the charge date/time, auto-calculate the 24 h expiration, and capture
Recorded By / Verified By signatures.

FM set (12, all in DE1 100):
  ZSMPL_FM_CUSTOM_INDEX          leftmost # row index                        (FMDATA)
  ZSMPL_FM_GET_EXPIRY_DATE       Batch No -> Exp. Date (validation)           (FMDATA)
  ZSMPL_FM_CREATE_PROC_MSG       goods issue Z_PICONS mvt 261 on signing      (FMDATA)
  /SMPL/PPPI_FM_GET_DATE_TIME    WFI Charge Date & Time stamp                 (biosamp)
  ZSMPL_FM_GET_EXP_DATE          WFI Expiration Date = charge date + 1 day    (FMDATA)
  /SMPL/PPPI_FM_SIG_ADD_DB_CB    Performed-By / Recorded-By signature         (golden)
  /SMPL/PPPI_FM_SIG_POPULATE_CB  signature callback                          (golden)
  /SMPL/PPPI_FM_SIG_VALIDATION   signature validator                         (golden)
  /SMPL/PPPI_FM_VALI_SUPE_SIG    Verified-By (verification) signature         (golden)
  /SMPL/PPPI_FM_INITIAL_ACTIVE   step activation                             (golden)
  /SMPL/MBR_DEP_ADD_PERFORM      Performed-By dependency                      (golden)
  /SMPL/MBR_DEP_CHECK_ACTIVE     active-check dependency                      (golden)

The Item column (Bag / Filter) is a restricted-value CT04 characteristic (configuration
via PPPI_REQUESTED_VALUE, e.g. ZSMPL_CHAR_WFI_ITEM) - no function module.
"""
import sys, io, copy, os, shutil, zipfile, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import build_phase2 as B
from p2lib import is_heading, heading_text

ROOT='..'
DONOR=ROOT+'/Phase 2 XSteps/SMPL- Three Variable Calc/SMPL_Three Variable Calc XStep Design Specification.docx'
OUTDIR=ROOT+'/AZ Phase 3 XStep Mock Ups (POC)/WFI Container Setup and Charge'
OUT=OUTDIR+'/SMPL_WFI Container Setup & Charge XStep Design Specification.docx'

FM_ORDER=['ZSMPL_FM_CUSTOM_INDEX','ZSMPL_FM_GET_EXPIRY_DATE','ZSMPL_FM_CREATE_PROC_MSG',
          '/SMPL/PPPI_FM_GET_DATE_TIME','ZSMPL_FM_GET_EXP_DATE',
          '/SMPL/PPPI_FM_SIG_ADD_DB_CB','/SMPL/PPPI_FM_SIG_POPULATE_CB','/SMPL/PPPI_FM_SIG_VALIDATION',
          '/SMPL/PPPI_FM_VALI_SUPE_SIG','/SMPL/PPPI_FM_INITIAL_ACTIVE','/SMPL/MBR_DEP_ADD_PERFORM',
          '/SMPL/MBR_DEP_CHECK_ACTIVE']

os.makedirs(OUTDIR, exist_ok=True)
shutil.copy(DONOR, OUT)
doc=Document(OUT); body=doc.element.body
children=list(body)
sectPr=children[-1] if children[-1].tag==qn('w:sectPr') else None
purpose=next(el for el in children if el.tag==qn('w:p') and is_heading(el) and heading_text(el).strip()=='Purpose')
anchor=copy.deepcopy(purpose)
start=children.index(purpose); end=children.index(sectPr) if sectPr is not None else len(children)
for el in children[start:end]: body.remove(el)

bid=B.max_bookmark_id(body)+1
def add(el):
    if sectPr is not None: sectPr.addprevious(el)
    else: body.append(el)
def H(t):
    global bid
    add(B.make_heading_like(anchor,t,bid)); bid+=1
def P(t): add(B.mk_plain(t))
def Bul(t): add(B.mk_bullet(t))
def BL(): add(B.mk_blank())
def TBL(tmpl, headers, rows):
    t=B.mk_table(tmpl, rows); hdr=t.findall(qn('w:tr'))[0]; tcs=hdr.findall(qn('w:tc'))
    for i,h in enumerate(headers): B.set_cell_text(tcs[i], h)
    add(t)
def T3(headers, rows): TBL(B.T_PTABLE, headers, rows)
def T2(headers, rows): TBL(B.T_ETABLE, headers, rows)
def IMG(path, width_in=6.5):
    para=doc.add_paragraph(); para.alignment=WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(path, width=Inches(width_in))
    el=para._p; el.getparent().remove(el); return el

# ---------------- content ----------------
H('Purpose')
P('The purpose of this document is to outline the design and configuration of the SMPL: WFI Container Setup & '
  'Charge XStep. This XStep documents the setup of the Water for Injection (WFI) collection vessel for a processing '
  'day: it records the collection bag and WFI filter used (with SAP goods-issue consumption), captures the WFI valve '
  'used, stamps the WFI charge date and time, automatically calculates the 24-hour WFI expiration, and captures the '
  'appropriate electronic signatures.')

H('Overview')
P('The SMPL: WFI Container Setup & Charge XStep provides the operator with a Data Input table and a set of charge '
  'fields. In the table the operator records the collection bag and the WFI filter as separate lines: the Item is '
  'chosen from a restricted-value dropdown (Bag / Filter); the Part No. is entered; the Batch No. is entered, which '
  'triggers an input validation that returns and displays the Exp. Date; and each line is signed Performed By. When '
  'a line is signed, a goods-issue process message (Z_PICONS, movement type 261 - order consumption) is raised to '
  'consume the material into SAP, so no separate "SAP Consumption" field or transaction is required. Rows are '
  'numbered automatically and can be added with Add Row.')
P('Below the table the operator records the WFI Valve ID, verifies the WFI port is flushed and charges WFI into the '
  'vessel, then presses the Record button to stamp the WFI Charge Date and Time. The WFI Expiration Date & Time is a '
  'read-only field calculated as 24 hours (one day) after the charge - the expiration date is computed from the '
  'charge date and the expiration time equals the charge time. The charge is completed with a Recorded-By signature '
  'and a Verified-By signature.')
P('The Exp. Date, the stamped Charge Date/Time and the calculated Expiration Date & Time are read-only outputs '
  'populated by function modules and are never hand-keyed. The Item dropdown is backed by a restricted-value CT04 '
  'characteristic. All function modules used by this XStep already exist in DE1 100; no new development is required.')

H('Reasons for developing')
P('The SMPL: WFI Container Setup & Charge XStep was developed to replace the manual WFI container setup and charge '
  'record in the paper batch record (Manufacturing Directions). Driving the Exp. Date from an input validation and '
  'the WFI expiration from a date calculation removes transcription and arithmetic error and guarantees a correct, '
  'tamper-evident 24-hour expiration; stamping the charge date/time with a system-time function module removes '
  'transcription error; raising the goods-issue process message guarantees SAP consumption of the bag and filter '
  'without a separate transaction; and capturing electronic signatures enforces the data-integrity and '
  'dual-verification requirements of the process.')

H('Authorization')
P('Access to the XStep and the ability to enter data and apply signatures is controlled by the standard SiMPL EBR '
  'security model. The operator must hold the PFCG role assigned to the relevant process order / control recipe. '
  'Performed-By, Recorded-By and Verified-By signatures are captured using the SAP digital signature method '
  'configured for the EBR.')

H('Assumptions/ Dependencies')
Bul('The generic SiMPL indexing, validation, date/time, expiry, goods-issue, signature and activation function '
    'modules (ZSMPL_FM_CUSTOM_INDEX, ZSMPL_FM_GET_EXPIRY_DATE, ZSMPL_FM_CREATE_PROC_MSG, '
    '/SMPL/PPPI_FM_GET_DATE_TIME, ZSMPL_FM_GET_EXP_DATE, /SMPL/PPPI_FM_SIG_ADD_DB_CB, '
    '/SMPL/PPPI_FM_SIG_POPULATE_CB, /SMPL/PPPI_FM_SIG_VALIDATION, /SMPL/PPPI_FM_VALI_SUPE_SIG, '
    '/SMPL/PPPI_FM_INITIAL_ACTIVE, /SMPL/MBR_DEP_ADD_PERFORM, /SMPL/MBR_DEP_CHECK_ACTIVE) exist and are active in '
    'the target system.')
Bul('A restricted-value CT04 characteristic (e.g. ZSMPL_CHAR_WFI_ITEM) with the values Bag and Filter exists and '
    'is assigned to the Item column (PPPI_REQUESTED_VALUE).')
Bul('The bag and WFI-filter parts exist in the material master and their batches exist in the batch master (MCH1) '
    'with a shelf-life expiration date; the collection vessel has been inspected and set up per SOP-0080854.')
Bul('The plant, storage location and movement type (261) are configured so the Z_PICONS goods-issue process '
    'message posts the order consumption for the bag and filter.')
Bul('The WFI hold time is 24 hours (one day); the WFI port is flushed before charging per the Manufacturing '
    'Directions.')

H('Validation Checks')
P('The following validations are applied within the XStep:')
T3(['Field','Validation','Function Module'],
   [['Item','Restricted to the Bag / Filter values of a CT04 characteristic','(CT04 characteristic - configuration)'],
    ['Part No.','Free-text entry of the bag / WFI-filter part number; required','(none)'],
    ['Batch No.','Valid material/batch combination; returns the Exp. Date; not expired','ZSMPL_FM_GET_EXPIRY_DATE'],
    ['Exp. Date','Read-only output populated from the Batch No. validation','ZSMPL_FM_GET_EXPIRY_DATE'],
    ['SAP consumption','Goods-issue process message (Z_PICONS, movement type 261) raised when a line is signed','ZSMPL_FM_CREATE_PROC_MSG'],
    ['Performed By','Mandatory electronic signature per row','/SMPL/PPPI_FM_SIG_ADD_DB_CB'],
    ['WFI Valve ID','Free-text entry; required','(none)'],
    ['WFI Charge Date & Time','System-stamped (read-only) when the Record button is pressed','/SMPL/PPPI_FM_GET_DATE_TIME'],
    ['WFI Expiration Date & Time','Read-only; expiration date = charge date + 1 day (24 h); time = charge time','ZSMPL_FM_GET_EXP_DATE'],
    ['Recorded By','Mandatory electronic signature for the WFI charge','/SMPL/PPPI_FM_SIG_ADD_DB_CB'],
    ['Verified By','Mandatory verification (supervisory) signature at step level','/SMPL/PPPI_FM_VALI_SUPE_SIG']])
BL()

H('XStep Layout Design')
P('The XStep is rendered as a Data Input table with an instruction panel and a set of charge fields below the '
  'table. The instruction reads: "Set up the Day 1 / Day 2 WFI collection vessel per SOP-0080854. Inspect the '
  'collection vessel, then record the bag and WFI filter as separate lines in the table below - select the Item '
  'from the dropdown (Bag / Filter - a restricted-value CT04 characteristic), enter each Part No. and Batch No. '
  '(the Exp. Date auto-populates via input validation) and sign Performed By per line. Each line is consumed via '
  'the Goods Issue process message (Z_PICONS, mvt 261). Below the table, record the WFI valve ID, verify the WFI '
  'port is flushed and charge WFI into the vessel; stamp the charge date/time (expiration is 24 hours from charge, '
  'auto-calculated)." Each table row captures the following:')
Bul('# - read-only row index, numbered automatically.')
Bul('Item - dropdown restricted to Bag / Filter (CT04 characteristic).')
Bul('Part No. - entry.')
Bul('Batch No. - entry; on entry an input validation returns the Exp. Date.')
Bul('Exp. Date - read-only, populated by the Batch No. validation.')
Bul('Performed By - signature, required (one per row); signing raises the Z_PICONS goods-issue consumption.')
P('The following fields are captured below the table (Add Row is available on the table):')
Bul('WFI Valve ID - entry, required.')
Bul('Record - button that stamps the WFI Charge Date and Time.')
Bul('WFI Charge Date & Time - read-only, system-stamped Date then Time.')
Bul('WFI Expiration Date & Time (Charge + 24 h) - read-only, calculated (charge date + 1 day at the charge time).')
Bul('Recorded By - signature, required.')
Bul('Verified By - verification (supervisory) signature, required.')
add(IMG(OUTDIR+'/image.png'))
BL()

H('Function Module(s)')
P('The following function modules are used by this XStep. All of them already exist in DE1 100 and are reused '
  'as-is; no new function module is required.')
for fm in FM_ORDER:
    for el in B.lib_fm_elems(fm): add(el)

H('Pseudocode')
for fm in FM_ORDER:
    for el in B.lib_pseudo_elems(fm): add(el)

H('Configuration Specifications')
P('The XStep is configured using the following parameters:')
T3(['Parameter','Value / Configuration','Description'],
   [['Header','WFI Container Setup & Charge','Header text displayed at the top of the step (configurable in SiMPL MBR).'],
    ['Instruction','(instruction text)','The instruction text displayed to the operator (configurable in SiMPL MBR).'],
    ['Item','Dropdown (CT04 characteristic)','Restricted to Bag / Filter via PPPI_REQUESTED_VALUE (e.g. ZSMPL_CHAR_WFI_ITEM).'],
    ['Part No.','Entry field','Free-text entry for the bag / WFI-filter part number.'],
    ['Batch No.','Entry field, input validation','Validated by ZSMPL_FM_GET_EXPIRY_DATE, which returns the Exp. Date.'],
    ['Exp. Date','Read-only output','Populated by the Batch No. validation (ZSMPL_FM_GET_EXPIRY_DATE).'],
    ['Goods Issue','Z_PICONS process message','Raised by ZSMPL_FM_CREATE_PROC_MSG on signing a line (movement type 261).'],
    ['Performed By','Signature column','Per-row Performed-By signature (/SMPL/PPPI_FM_SIG_ADD_DB_CB).'],
    ['WFI Valve ID','Entry field','Free-text entry for the WFI valve identifier (required).'],
    ['Record button','Stamp WFI Charge Date/Time','Calls /SMPL/PPPI_FM_GET_DATE_TIME to populate the Charge Date and Time.'],
    ['WFI Expiration Date & Time','Read-only, calculated','ZSMPL_FM_GET_EXP_DATE with days = 1 (charge date + 24 h); time = charge time.'],
    ['Recorded By','Signature field','Recorded-By signature for the WFI charge (/SMPL/PPPI_FM_SIG_ADD_DB_CB).'],
    ['Verified By','Step-level signature','Verification/supervisory signature (/SMPL/PPPI_FM_VALI_SUPE_SIG).'],
    ['Add Row','Enabled','Rows numbered automatically by ZSMPL_FM_CUSTOM_INDEX.']])
BL()

H('Test Scenarios')
T3(['#','Scenario','Expected Result'],
   [['1','Select the Item dropdown on a row.','Only the restricted values Bag and Filter are available for selection.'],
    ['2','Enter a valid Batch No. for the entered Part No.','The Exp. Date is returned and displayed read-only (ZSMPL_FM_GET_EXPIRY_DATE).'],
    ['3','Enter a Batch No. that does not exist for the material.','The entry is rejected with BATCH_COMBO_NOT_FOUND.'],
    ['4','Enter a Batch No. whose expiration date is in the past.','The entry is rejected with EXPIRY_DATE_IN_PAST.'],
    ['5','Apply the Performed-By signature on a completed row.','The row is signed and a Z_PICONS goods-issue process message (movement type 261) is raised to consume the material into SAP (ZSMPL_FM_CREATE_PROC_MSG).'],
    ['6','Add a second row with Add Row.','A new numbered row is added and can be completed independently.'],
    ['7','Press the Record button after charging WFI.','The WFI Charge Date and Time are stamped with the current system date and time (read-only).'],
    ['8','Observe the WFI Expiration Date & Time after the charge is stamped.','The expiration is calculated as 24 hours after the charge (charge date + 1 day at the charge time) and displayed read-only (ZSMPL_FM_GET_EXP_DATE).'],
    ['9','Enter the WFI Valve ID and attempt to complete the step without the Recorded-By or Verified-By signature.','The step cannot be completed; both signatures are mandatory.'],
    ['10','Re-open the completed step.','The recorded values, validation outputs, stamped/calculated dates and signatures are displayed and retained.']])
BL()

H('Document References')
T2(['Reference No.','Document Title'],
   [['1','AZD8630 - Affinity (CaptureSelect CH1-XL) Manufacturing Directions (MABR-0023643, PN 8010457)'],
    ['2','SOP-0080854 - WFI Collection Vessel Setup'],
    ['3','SiMPL XStep Library - Process Manufacturing (DE1 100)'],
    ['4','SMPL: Room/Equipment Assign XStep Design Specification V1.3 (reference template)']])
BL()

H('Revision History')
T3(['Version No.','Description','Date'],
   [['1.0','Initial document','2026-07-14']])

# ---------------- fixes + save ----------------
B.neutralise_empty_headings(body)
B.set_update_fields(doc)
doc.save(OUT)

# title-canvas + header/footer text: rename the XStep (XML-escape the ampersand)
tmp=OUT+'.tmp'
with zipfile.ZipFile(OUT) as z: data={n:z.read(n) for n in z.namelist()}
for n in list(data):
    if re.match(r'word/(document|header\d+|footer\d+)\.xml$', n):
        x=data[n].decode('utf-8')
        x2=x.replace('Three Variable Calc','WFI Container Setup &amp; Charge')
        if x2!=x: data[n]=x2.encode('utf-8')
with zipfile.ZipFile(tmp,'w',zipfile.ZIP_DEFLATED) as z:
    for n,b in data.items(): z.writestr(n,b)
shutil.move(tmp,OUT)
print('built',OUT)
