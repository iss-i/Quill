# -*- coding: utf-8 -*-
"""Build the SMPL: Additional Manufacturing Supplies XStep Design Specification (.docx)
in the AZ template. From-scratch build (AZ Phase 3 has only a mock-up): clones the
Three Variable Calc doc for the shell (title page, styles, live TOC), rebuilds the 13
sections from the Additional Manufacturing Supplies mock-up, and reuses the Phase-2
FM/pseudocode emitters. This XStep is a VARIANT of SMPL: Additional Assembly (+ Material
Consumption) driven by input validations instead of a Get-Materials button.

FM set (12, all in DE1 100):
  ZSMPL_FM_CUSTOM_INDEX          leftmost # row index                       (FMDATA)
  ZSMPL_FM_GET_MAT_DETAIL        Part No -> Material Description (validation) (authored, local src)
  ZSMPL_FM_GET_EXPIRY_DATE       Batch No -> Exp. Date (validation)          (authored, DE1_100 src)
  ZSMPL_FM_CHECK_CHAR_DATE       Autoclave Exp. Date validator               (authored, DE1_100 src)
  ZSMPL_FM_CREATE_PROC_MSG       goods issue Z_PICONS mvt 261 on signing     (authored, DE1_100 src)
  /SMPL/PPPI_FM_SIG_ADD_DB_CB    Performed-By signature (per row)            (golden)
  /SMPL/PPPI_FM_SIG_POPULATE_CB  signature callback                         (golden)
  /SMPL/PPPI_FM_SIG_VALIDATION   signature validator                        (golden)
  /SMPL/PPPI_FM_VALI_SUPE_SIG    Witnessed-By signature (step level)         (golden)
  /SMPL/PPPI_FM_INITIAL_ACTIVE   step activation                            (golden)
  /SMPL/MBR_DEP_ADD_PERFORM      Performed-By dependency                     (golden)
  /SMPL/MBR_DEP_CHECK_ACTIVE     active-check dependency                     (golden)
"""
import sys, io, copy, os, shutil, zipfile, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import build_phase2 as B
from p2lib import is_heading, heading_text

ROOT='..'
DONOR=ROOT+'/Phase 2 XSteps/SMPL- Three Variable Calc/SMPL_Three Variable Calc XStep Design Specification.docx'
OUTDIR=ROOT+'/AZ Phase 3 XStep Mock Ups/Additional Manufacturing Supplies'
OUT=OUTDIR+'/SMPL_Additional Manufacturing Supplies XStep Design Specification.docx'

# ---------------------------------------------------------------------------
# Authored FMs (sourced from DE1 100 / local Function Modules folder), injected
# into the FM library so the emitters render them in the accepted style.
# ---------------------------------------------------------------------------
B.FMDATA['ZSMPL_FM_GET_MAT_DETAIL']={
 'name':'ZSMPL_FM_GET_MAT_DETAIL',
 'imports':[('IM_MATNR','MATNR','Material (part) number entered by the operator')],
 'exports':[('EX_MAT_DESC','MAKTX','Material description returned for the part number'),
            ('EX_UOM','MEINS','Base unit of measure of the material')],
 'exceptions':[('MAT_DESC_DNE_FOR_LANGU','No material description exists for the logon language'),
               ('UOM_NOT_FOUND','No base unit of measure found for the material')],
 'pseudo':[
   ('h','1. Purpose'),
   ('p','Retrieves the material description and base unit of measure for a material (part) number so the '
        'Material Description column is populated automatically when the operator enters a Part No.'),
   ('h','2. Input Parameters'),
   ('b','IM_MATNR - the material (part) number entered by the operator.'),
   ('h','3. Processing Logic'),
   ('s','a. Move IM_MATNR to a working material variable.'),
   ('s','b. Select the material description (MAKTX) from table MAKT for the material in the logon language (SY-LANGU).'),
   ('s','c. If no description is found in the logon language, re-select MAKT for the material in English ("E").'),
   ('s','d. Select the base unit of measure (MEINS) from table MARA for the material.'),
   ('s','e. If no unit of measure is found, raise UOM_NOT_FOUND.'),
   ('s','f. Return the description in EX_MAT_DESC and the unit of measure in EX_UOM.'),
   ('h','Additional Notes'),
   ('h','Error Handling:'),
   ('b','MAT_DESC_DNE_FOR_LANGU - no material description exists for the material in the requested language.'),
   ('b','UOM_NOT_FOUND - the material has no base unit of measure in MARA.'),
   ('h','Dependencies:'),
   ('b','Material master tables MAKT (descriptions) and MARA (base data) must be maintained for the part.'),
   ('h','Assumptions:'),
   ('b','The part number entered corresponds to a material master record; English is available as a fallback language.'),
 ],
}
B.FMDATA['ZSMPL_FM_GET_EXPIRY_DATE']={
 'name':'ZSMPL_FM_GET_EXPIRY_DATE',
 'imports':[('IM_MATNR','MATNR','Material (part) number of the supply (optional)'),
            ('IM_BATCH','CHARG_D','Batch number entered by the operator (optional)')],
 'exports':[('EX_EXP_DATE','VFDAT','Expiration (shelf-life) date returned for the material/batch')],
 'exceptions':[('MAT_IS_INITIAL','No material was supplied'),
               ('BATCH_COMBO_NOT_FOUND','The material/batch combination does not exist in SAP'),
               ('EXPIRY_DATE_IN_PAST','The batch expiration date is earlier than the current date')],
 'pseudo':[
   ('h','1. Purpose'),
   ('p','Validates the batch entered against the material and returns the batch expiration date so the Exp. Date '
        'column is populated automatically when the operator enters a Batch No.'),
   ('h','2. Input Parameters'),
   ('b','IM_MATNR - the material (part) number of the supply.'),
   ('b','IM_BATCH - the batch number entered by the operator.'),
   ('h','3. Processing Logic'),
   ('s','a. If IM_MATNR is blank, raise MAT_IS_INITIAL.'),
   ('s','b. Select the expiration date (VFDAT) from batch table MCH1 for the material and batch.'),
   ('s','c. If the material/batch combination is not found, raise BATCH_COMBO_NOT_FOUND.'),
   ('s','d. If the expiration date is earlier than the current date (SY-DATLO), raise EXPIRY_DATE_IN_PAST.'),
   ('s','e. Return the expiration date in EX_EXP_DATE.'),
   ('h','Additional Notes'),
   ('h','Error Handling:'),
   ('b','MAT_IS_INITIAL - the material number was not supplied.'),
   ('b','BATCH_COMBO_NOT_FOUND - no MCH1 record exists for the material/batch combination.'),
   ('b','EXPIRY_DATE_IN_PAST - the batch expiration date has already passed.'),
   ('h','Dependencies:'),
   ('b','Batch master table MCH1 must hold the material/batch with a shelf-life expiration date (VFDAT).'),
   ('h','Assumptions:'),
   ('b','The part number has already been validated (via ZSMPL_FM_GET_MAT_DETAIL); batches are managed with '
       'shelf-life expiration in SAP.'),
 ],
}
B.FMDATA['ZSMPL_FM_CHECK_CHAR_DATE']={
 'name':'ZSMPL_FM_CHECK_CHAR_DATE',
 'imports':[
   ('IV_VALUE','CHAR10','Date value entered by the operator, in yyyy.mm.dd format (or "N/A") (optional)'),
   ('IV_DATE','DATS','Date value supplied directly in internal (DATS) format (optional)'),
 ],
 'exceptions':[
   ('INCORRECT_DATE_FORMAT','Entered value is not in the expected yyyy.mm.dd format'),
   ('INCORRECT_VALUE_INPUT','Entered value is neither a valid date nor "N/A"'),
   ('DATE_CANNOT_BE_IN_PAST','The entered/derived date is earlier than the current system date'),
 ],
 'pseudo':[
   ('h','1. Purpose'),
   ('p','Validates a character-format date entered by the operator (for example the Autoclave Exp. Date). '
        'Confirms the value is either a valid date in yyyy.mm.dd format or the literal "N/A", and that any date '
        'supplied is not earlier than the current system date.'),
   ('h','2. Input Parameters'),
   ('b','IV_VALUE - the date the operator typed, as a 10-character string in yyyy.mm.dd format, or "N/A" / "NA" (optional).'),
   ('b','IV_DATE - an alternative date supplied directly in internal (DATS) format (optional).'),
   ('h','3. Processing Logic'),
   ('s','a. Take a working copy of IV_VALUE, translate it to upper case and measure its length; build an internal '
        'date value by removing the "." separators.'),
   ('s','b. If IV_VALUE is not blank:'),
   ('b','If it equals "N/A" or "NA", accept the entry and perform no further checks.'),
   ('b','Otherwise, if the month portion (positions 6-7) is greater than 12, raise INCORRECT_DATE_FORMAT with the '
        'message "Incorrect date format entered - use yyyy.mm.dd".'),
   ('b','Otherwise, if the value is shorter than 10 characters, raise INCORRECT_VALUE_INPUT with the message '
        '"Incorrect value input - options are Date or N/A".'),
   ('b','Otherwise, if the resulting date is earlier than the system date, raise DATE_CANNOT_BE_IN_PAST.'),
   ('s','c. If IV_DATE is supplied and is earlier than the system date, raise DATE_CANNOT_BE_IN_PAST.'),
   ('h','Additional Notes'),
   ('h','Error Handling:'),
   ('b','INCORRECT_DATE_FORMAT - the month component is invalid / the value is not in yyyy.mm.dd format.'),
   ('b','INCORRECT_VALUE_INPUT - the value is neither a full date nor "N/A".'),
   ('b','DATE_CANNOT_BE_IN_PAST - the entered or derived date precedes the current system date.'),
   ('h','Dependencies:'),
   ('b','None - pure ABAP date handling; no database reads or external function-module calls.'),
   ('h','Assumptions:'),
   ('b','Operators enter dates in yyyy.mm.dd format; "N/A" is permitted where an autoclave expiry does not apply.'),
 ],
}
B.FMDATA['ZSMPL_FM_CREATE_PROC_MSG']={
 'name':'ZSMPL_FM_CREATE_PROC_MSG',
 'imports':[
   ('I_PLANT','WERKS_D','Plant for the goods issue'),
   ('I_DATE','DATE','Event date; defaults to the current date if blank (optional)'),
   ('I_TIME','TIMS','Event time; defaults to the current time if blank (optional)'),
   ('I_MATNR','MATNR','Material (part) number to be consumed'),
   ('I_GI_QTY','ERFMG','Goods-issue (consumed) quantity'),
   ('I_ORDER','AUFNR','Process order (optional)'),
   ('I_ERFME','MEINS','Unit of measure (optional)'),
   ('I_BWART','BWART','Movement type - 261 for order consumption (optional)'),
   ('I_RSNUM','RSNUM','Reservation number (optional)'),
   ('I_RSPOS','RSPOS','Reservation item (optional)'),
   ('I_MSG_HEAD','CO_SOURCE','Process-message sender name; defaults to the user (optional)'),
   ('I_BATCH','CHARG_D','Batch of the material (optional)'),
   ('I_PHASE','VORNR','Phase / operation number (optional)'),
   ('I_STLOC','LGORT_D','Storage location for the goods issue'),
   ('I_FLAG','CHAR1','Conditional trigger flag (optional)'),
   ('I_MATCH_FLAG','CHAR1','Value the trigger flag must match to proceed (optional)'),
 ],
 'exceptions':[('MESSAGE_CREATION_FAIL','The process message could not be created'),
               ('CHARACTERISTIC_ERROR','A process-message characteristic was rejected')],
 'pseudo':[
   ('h','1. Purpose'),
   ('p','Creates the custom goods-issue process message Z_PICONS to consume a material into SAP (movement type '
        '261, order consumption) when the operator signs a supply line, replacing a manual SAP goods-issue '
        'transaction.'),
   ('h','2. Input Parameters'),
   ('b','I_PLANT / I_STLOC - plant and storage location for the goods issue.'),
   ('b','I_MATNR / I_BATCH / I_GI_QTY / I_ERFME - the material, batch, consumed quantity and unit of measure.'),
   ('b','I_ORDER / I_PHASE / I_RSNUM / I_RSPOS - process order, phase and reservation reference.'),
   ('b','I_BWART - movement type (261 for order consumption).'),
   ('b','I_DATE / I_TIME - event date and time (default to the system date/time).'),
   ('b','I_FLAG / I_MATCH_FLAG - optional conditional trigger.'),
   ('h','3. Processing Logic'),
   ('s','a. If I_FLAG is set and does not equal I_MATCH_FLAG, return without creating a message.'),
   ('s','b. Condense and upper-case the material and batch; if either is "N/A"/"NA", or the material is blank, '
        'return without creating a message.'),
   ('s','c. Default the event date and time to the system date/time when not supplied.'),
   ('s','d. Build the process-message header for category Z_PICONS (plant and sender name).'),
   ('s','e. Build the characteristics: event date/time, material, quantity consumed, process order, unit of '
        'measure, movement type, reservation and item, batch, phase and storage location.'),
   ('s','f. Call BAPI_PROCESS_MESSAGE_CREATEMLT to create the message.'),
   ('s','g. Inspect the header, characteristic and general return tables; raise CHARACTERISTIC_ERROR or '
        'MESSAGE_CREATION_FAIL on any error message.'),
   ('h','Additional Notes'),
   ('h','Error Handling:'),
   ('b','CHARACTERISTIC_ERROR - a characteristic value was rejected by the process-message BAPI.'),
   ('b','MESSAGE_CREATION_FAIL - the header or general return contained an error message.'),
   ('b','Materials/batches of "N/A" and blank materials are skipped silently (no message and no error).'),
   ('h','Dependencies:'),
   ('b','Process message category Z_PICONS and the standard BAPI_PROCESS_MESSAGE_CREATEMLT; the message is '
       'posted by the process-message destination that executes the movement.'),
   ('h','Assumptions:'),
   ('b','The plant, storage location and movement type are configured for order consumption; the material and '
       'batch have been validated on the line before signing.'),
 ],
}

FM_ORDER=['ZSMPL_FM_CUSTOM_INDEX','ZSMPL_FM_GET_MAT_DETAIL','ZSMPL_FM_GET_EXPIRY_DATE',
          'ZSMPL_FM_CHECK_CHAR_DATE','ZSMPL_FM_CREATE_PROC_MSG',
          '/SMPL/PPPI_FM_SIG_ADD_DB_CB','/SMPL/PPPI_FM_SIG_POPULATE_CB','/SMPL/PPPI_FM_SIG_VALIDATION',
          '/SMPL/PPPI_FM_VALI_SUPE_SIG','/SMPL/PPPI_FM_INITIAL_ACTIVE','/SMPL/MBR_DEP_ADD_PERFORM',
          '/SMPL/MBR_DEP_CHECK_ACTIVE']

os.makedirs(OUTDIR, exist_ok=True)
shutil.copy(DONOR, OUT)
doc=Document(OUT); body=doc.element.body
children=list(body)
sectPr=children[-1] if children[-1].tag==qn('w:sectPr') else None
purpose=next(el for el in children if el.tag==qn('w:p') and is_heading(el) and heading_text(el).strip()=='Purpose')
anchor=copy.deepcopy(purpose)
start=children.index(purpose); end=children.index(sectPr) if sectPr is not None else len(children)
for el in children[start:end]: body.remove(el)

bid=B.max_bookmark_id(body)+1
def add(el):
    if sectPr is not None: sectPr.addprevious(el)
    else: body.append(el)
def H(t):
    global bid
    add(B.make_heading_like(anchor,t,bid)); bid+=1
def P(t): add(B.mk_plain(t))
def Bul(t): add(B.mk_bullet(t))
def BL(): add(B.mk_blank())
def TBL(tmpl, headers, rows):
    t=B.mk_table(tmpl, rows); hdr=t.findall(qn('w:tr'))[0]; tcs=hdr.findall(qn('w:tc'))
    for i,h in enumerate(headers): B.set_cell_text(tcs[i], h)
    add(t)
def T3(headers, rows): TBL(B.T_PTABLE, headers, rows)
def T2(headers, rows): TBL(B.T_ETABLE, headers, rows)
def IMG(path, width_in=6.5):
    para=doc.add_paragraph(); para.alignment=WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(path, width=Inches(width_in))
    el=para._p; el.getparent().remove(el); return el

# ---------------- content ----------------
H('Purpose')
P('The purpose of this document is to outline the design and configuration of the SMPL: Additional Manufacturing '
  'Supplies XStep. This XStep records any additional manufacturing supplies (filters, bags, assemblies and similar '
  'items) used during the process, capturing the applicable MPR step, the part number and its material description, '
  'the batch number and expiration date, the serial number, and, where applicable, the autoclave assembly ID and '
  'autoclave expiration date, with automatic SAP goods-issue consumption and the appropriate electronic signatures.')

H('Overview')
P('The SMPL: Additional Manufacturing Supplies XStep provides the operator with a Data Input table that documents '
  'each additional supply used during processing. For each item the operator records the MPR Step No.; enters the '
  'Part No., which triggers an input validation that returns and displays the Material Description; enters the Batch '
  'No., which triggers an input validation that returns and displays the Exp. Date; and records the Serial No. and, '
  'where applicable, the Autoclave Assembly ID and Autoclave Exp. Date. Each row is completed with a Performed-By '
  'signature, and the step is closed with a Witness signature.')
P('The Material Description and Exp. Date are read-only outputs populated by the input-validation function modules '
  'when the Part No. and Batch No. are entered; they are never hand-keyed. The Autoclave Exp. Date is validated as a '
  'date. When a supply line is signed, a goods-issue process message (Z_PICONS, movement type 261 - order '
  'consumption) is raised to consume the material into SAP, so no separate "SAP Consumption" field or transaction is '
  'required. Additional rows can be added with the Add Row action, and the rows are numbered automatically.')
P('This XStep is a variant of the SMPL: Additional Assembly / SMPL: Material Consumption building blocks, configured '
  'to drive the material description and expiry from input validations rather than a Get-Materials button. All '
  'function modules used by this XStep already exist in DE1 100; no new development is required.')

H('Reasons for developing')
P('The SMPL: Additional Manufacturing Supplies XStep was developed to replace the manual "Additional Manufacturing '
  'Supplies Used in Process" record in the paper batch record (Manufacturing Directions). Driving the Material '
  'Description and Exp. Date from input validations removes transcription error and enforces that only valid part '
  'and batch combinations are recorded; raising the goods-issue process message guarantees SAP consumption without a '
  'separate manual transaction; and capturing electronic Performed-By and Witness signatures enforces the '
  'data-integrity and dual-verification requirements of the process. A single reusable XStep replaces a supplies '
  'table that would otherwise be reprinted for every occurrence.')

H('Authorization')
P('Access to the XStep and the ability to enter data and apply signatures is controlled by the standard SiMPL EBR '
  'security model. The operator must hold the PFCG role assigned to the relevant process order / control recipe. '
  'Performed-By and Witness signatures are captured using the SAP digital signature method configured for the EBR.')

H('Assumptions/ Dependencies')
Bul('The generic SiMPL indexing, validation, goods-issue, signature and activation function modules '
    '(ZSMPL_FM_CUSTOM_INDEX, ZSMPL_FM_GET_MAT_DETAIL, ZSMPL_FM_GET_EXPIRY_DATE, ZSMPL_FM_CHECK_CHAR_DATE, '
    'ZSMPL_FM_CREATE_PROC_MSG, /SMPL/PPPI_FM_SIG_ADD_DB_CB, /SMPL/PPPI_FM_SIG_POPULATE_CB, '
    '/SMPL/PPPI_FM_SIG_VALIDATION, /SMPL/PPPI_FM_VALI_SUPE_SIG, /SMPL/PPPI_FM_INITIAL_ACTIVE, '
    '/SMPL/MBR_DEP_ADD_PERFORM, /SMPL/MBR_DEP_CHECK_ACTIVE) exist and are active in the target system.')
Bul('The parts entered exist in the material master (MAKT/MARA) and the batches exist in the batch master (MCH1) '
    'with a shelf-life expiration date; the material has been verified as acceptable per Section 4.')
Bul('The plant, storage location and movement type (261) are configured so the Z_PICONS goods-issue process '
    'message posts the order consumption for the material.')
Bul('The Autoclave Exp. Date is recorded in yyyy.mm.dd format, or "N/A" where an autoclave expiry does not apply.')
Bul('The reason for using an additional supply is documented in the Comments Section.')

H('Validation Checks')
P('The following validations are applied within the XStep:')
T3(['Field','Validation','Function Module'],
   [['MPR Step No.','Free-text entry of the applicable MPR step; required','(none)'],
    ['Part No.','Must be a valid SAP material; returns the Material Description','ZSMPL_FM_GET_MAT_DETAIL'],
    ['Material Description','Read-only output populated from the Part No. validation','ZSMPL_FM_GET_MAT_DETAIL'],
    ['Batch No.','Must be a valid material/batch combination; returns the Exp. Date; not expired','ZSMPL_FM_GET_EXPIRY_DATE'],
    ['Exp. Date','Read-only output populated from the Batch No. validation','ZSMPL_FM_GET_EXPIRY_DATE'],
    ['Serial No.','Free-text entry','(none)'],
    ['Autoclave Assembly ID','Free-text entry','(none)'],
    ['Autoclave Exp. Date','Must be a valid date in yyyy.mm.dd format (or "N/A") and not in the past','ZSMPL_FM_CHECK_CHAR_DATE'],
    ['SAP consumption','Goods-issue process message (Z_PICONS, movement type 261) raised when a line is signed','ZSMPL_FM_CREATE_PROC_MSG'],
    ['Performed By','Mandatory electronic signature per row','/SMPL/PPPI_FM_SIG_ADD_DB_CB'],
    ['Witnessed By','Mandatory witness (supervisory) signature at step level','/SMPL/PPPI_FM_VALI_SUPE_SIG']])
BL()

H('XStep Layout Design')
P('The XStep is rendered as a Data Input table with an instruction panel. The instruction reads: "If additional '
  'materials (filter, bag, assembly, etc.) are used during processing, verify the material is acceptable per '
  'Section 4 and record its information; document the reason for the change in the Comments Section. Enter the Part '
  'No. - an input validation displays the Material Description. Enter the Batch No. - the Exp. Date is '
  'auto-populated. Serial No. and Autoclave Assembly ID / Exp. Date are recorded manually. SAP consumption is posted '
  'automatically by the Goods Issue process message (Z_PICONS, movement type 261) - no manual SAP Consumption field '
  'is required." Each table row captures the following:')
Bul('# - read-only row index, numbered automatically.')
Bul('MPR Step No. - entry, the applicable step in the Manufacturing Directions.')
Bul('Part No. - entry; on entry an input validation returns the Material Description.')
Bul('Material Description - read-only, populated by the Part No. validation.')
Bul('Batch No. - entry; on entry an input validation returns the Exp. Date.')
Bul('Exp. Date - read-only, populated by the Batch No. validation.')
Bul('Serial No. - entry.')
Bul('Autoclave Assembly ID - entry.')
Bul('Autoclave Exp. Date - entry, validated as a date (yyyy.mm.dd or "N/A").')
Bul('Performed By - signature, required (one per row); signing raises the Z_PICONS goods-issue consumption.')
P('Additional rows can be added with Add Row; the rows are numbered automatically. A step-level Witnessed By '
  'signature is captured in the footer.')
add(IMG(OUTDIR+'/image.png'))
BL()

H('Function Module(s)')
P('The following function modules are used by this XStep. All of them already exist in DE1 100 and are reused '
  'as-is; no new function module is required.')
for fm in FM_ORDER:
    for el in B.lib_fm_elems(fm): add(el)

H('Pseudocode')
for fm in FM_ORDER:
    for el in B.lib_pseudo_elems(fm): add(el)

H('Configuration Specifications')
P('The XStep is configured using the following parameters:')
T3(['Parameter','Value / Configuration','Description'],
   [['Header','Additional Manufacturing Supplies','Header text displayed at the top of the step (configurable in SiMPL MBR).'],
    ['Instruction','(instruction text)','The instruction text displayed to the operator (configurable in SiMPL MBR).'],
    ['MPR Step No.','Entry field','Free-text entry for the applicable Manufacturing Directions step (required).'],
    ['Part No.','Entry field, input validation','Validated by ZSMPL_FM_GET_MAT_DETAIL, which returns the Material Description.'],
    ['Material Description','Read-only output','Populated by the Part No. validation (ZSMPL_FM_GET_MAT_DETAIL).'],
    ['Batch No.','Entry field, input validation','Validated by ZSMPL_FM_GET_EXPIRY_DATE, which returns the Exp. Date.'],
    ['Exp. Date','Read-only output','Populated by the Batch No. validation (ZSMPL_FM_GET_EXPIRY_DATE).'],
    ['Serial No.','Entry field','Free-text entry for the item serial number.'],
    ['Autoclave Assembly ID','Entry field','Free-text entry for the autoclave assembly identifier.'],
    ['Autoclave Exp. Date','Entry field, date validator','Validated by ZSMPL_FM_CHECK_CHAR_DATE (yyyy.mm.dd or "N/A", not in past).'],
    ['Goods Issue','Z_PICONS process message','Raised by ZSMPL_FM_CREATE_PROC_MSG on signing (movement type 261, order consumption).'],
    ['Performed By','Signature column','Per-row Performed-By signature (/SMPL/PPPI_FM_SIG_ADD_DB_CB).'],
    ['Witnessed By','Step-level signature','Witness/supervisory signature (/SMPL/PPPI_FM_VALI_SUPE_SIG).'],
    ['Add Row','Enabled','Rows numbered automatically by ZSMPL_FM_CUSTOM_INDEX.']])
BL()

H('Test Scenarios')
T3(['#','Scenario','Expected Result'],
   [['1','Enter a valid Part No. for an existing material.','The Material Description is returned and displayed read-only (ZSMPL_FM_GET_MAT_DETAIL).'],
    ['2','Enter a Part No. that has no base unit of measure / material record.','The lookup fails (UOM_NOT_FOUND); the Material Description is not populated.'],
    ['3','Enter a valid Batch No. for the entered material.','The Exp. Date is returned and displayed read-only (ZSMPL_FM_GET_EXPIRY_DATE).'],
    ['4','Enter a Batch No. that does not exist for the material.','The entry is rejected with BATCH_COMBO_NOT_FOUND.'],
    ['5','Enter a Batch No. whose expiration date is in the past.','The entry is rejected with EXPIRY_DATE_IN_PAST.'],
    ['6','Enter an Autoclave Exp. Date with an invalid month (e.g. 2027.13.01).','The entry is rejected with "Incorrect date format entered - use yyyy.mm.dd" (INCORRECT_DATE_FORMAT).'],
    ['7','Enter "N/A" for the Autoclave Exp. Date.','The entry is accepted (no autoclave expiry applies).'],
    ['8','Apply the Performed-By signature on a completed row.','The row is signed and a Z_PICONS goods-issue process message (movement type 261) is raised to consume the material into SAP (ZSMPL_FM_CREATE_PROC_MSG).'],
    ['9','Add a second row with Add Row.','A new numbered row is added and can be completed independently.'],
    ['10','Attempt to complete the step without a Performed-By or Witness signature.','The step cannot be completed; both signatures are mandatory.'],
    ['11','Re-open the completed step.','The recorded values, validation outputs and signatures are displayed and retained.']])
BL()

H('Document References')
T2(['Reference No.','Document Title'],
   [['1','AZD8630 - Affinity (CaptureSelect CH1-XL) Manufacturing Directions (MABR-0023643, PN 8010457)'],
    ['2','SiMPL XStep Library - Process Manufacturing (DE1 100)'],
    ['3','SMPL: Room/Equipment Assign XStep Design Specification V1.3 (reference template)']])
BL()

H('Revision History')
T3(['Version No.','Description','Date'],
   [['1.0','Initial document','2026-07-14']])

# ---------------- fixes + save ----------------
B.neutralise_empty_headings(body)
B.set_update_fields(doc)
doc.save(OUT)

# title-canvas + header/footer text: rename the XStep
tmp=OUT+'.tmp'
with zipfile.ZipFile(OUT) as z: data={n:z.read(n) for n in z.namelist()}
for n in list(data):
    if re.match(r'word/(document|header\d+|footer\d+)\.xml$', n):
        x=data[n].decode('utf-8')
        x2=x.replace('Three Variable Calc','Additional Manufacturing Supplies')
        if x2!=x: data[n]=x2.encode('utf-8')
with zipfile.ZipFile(tmp,'w',zipfile.ZIP_DEFLATED) as z:
    for n,b in data.items(): z.writestr(n,b)
shutil.move(tmp,OUT)
print('built',OUT)
