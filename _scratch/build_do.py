# -*- coding: utf-8 -*-
"""Build the SMPL: Daily Observations XStep Design Specification (.docx) in the AZ
Phase 2 template. Daily Observations is the viral variant of Bioreactor Sampling
(adds CPE %, Cell Count Live/Dead, O2 Consumption). Clones the Three Variable Calc
doc for the shell (title page, styles, live TOC), rebuilds the body, and transplants
the 11 Bioreactor-Sampling FM blocks verbatim (all already exist in DE1 100)."""
import sys, io, copy, os, shutil, zipfile, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import build_phase2 as B
from p2lib import is_heading, heading_text

ROOT='..'
DONOR=ROOT+'/Phase 2 XSteps/SMPL- Three Variable Calc/SMPL_Three Variable Calc XStep Design Specification.docx'
OUTDIR=ROOT+'/Merck XStep Mock Ups/Daily Observations'
OUT=OUTDIR+'/SMPL_Daily Observations XStep Design Specification.docx'

# 8 FMs transplanted from the Bioreactor Sampling doc (they already exist in DE1 100).
# EXCLUDED: STORE_DAILY_SAMPLE (fixed glucose/BRX/VCD/viability interface doesn't fit the viral
# params - see Overview); GET_ASSIGNED_EQUI_EBR + ELB_FM_GET_ASS_EQ_VALID (the probes are recorded
# as entries, not order-assigned equipment); OPS (no probe-vs-external pH tolerance - pH is instead
# validated as a real value 0.0-14.0 via MIN_MAX).
FM_ORDER=['/SMPL/PPPI_FM_MIN_MAX','/SMPL/PPPI_FM_GET_DATE_TIME','/SMPL/PPPI_FM_SIG_ADD_DB_CB',
          '/SMPL/PPPI_FM_SIG_POPULATE_CB','/SMPL/PPPI_FM_SIG_VALIDATION','/SMPL/PPPI_FM_VALI_SUPE_SIG',
          '/SMPL/PPPI_FM_INITIAL_ACTIVE','/SMPL/MBR_DEP_CHECK_ACTIVE']

os.makedirs(OUTDIR, exist_ok=True)
shutil.copy(DONOR, OUT)
doc=Document(OUT); body=doc.element.body
children=list(body)
sectPr=children[-1] if children[-1].tag==qn('w:sectPr') else None
purpose=next(el for el in children if el.tag==qn('w:p') and is_heading(el) and heading_text(el).strip()=='Purpose')
anchor=copy.deepcopy(purpose)
start=children.index(purpose); end=children.index(sectPr) if sectPr is not None else len(children)
for el in children[start:end]: body.remove(el)

bid=B.max_bookmark_id(body)+1
def add(el):
    if sectPr is not None: sectPr.addprevious(el)
    else: body.append(el)
def H(t):
    global bid
    add(B.make_heading_like(anchor,t,bid)); bid+=1
def P(t): add(B.mk_plain(t))
def Bul(t): add(B.mk_bullet(t))
def BL(): add(B.mk_blank())
def TBL(tmpl, headers, rows):
    t=B.mk_table(tmpl, rows); hdr=t.findall(qn('w:tr'))[0]; tcs=hdr.findall(qn('w:tc'))
    for i,h in enumerate(headers): B.set_cell_text(tcs[i], h)
    add(t)
def T3(headers, rows): TBL(B.T_PTABLE, headers, rows)
def T2(headers, rows): TBL(B.T_ETABLE, headers, rows)
def IMG(path, width_in=6.5):
    para=doc.add_paragraph(); para.alignment=WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(path, width=Inches(width_in))
    el=para._p; el.getparent().remove(el); return el

# ---------------- content ----------------
H('Purpose')
P('The purpose of this document is to outline the design and configuration of the SMPL: Daily Observations '
  'XStep. This XStep records the daily in-process conditions of the production bioreactor - temperature, '
  'agitation, pH, dissolved oxygen, oxygen consumption, cytopathic effect and live/dead cell counts - sampled '
  'from the reactor each day, with a system-stamped sample date and time, in-range validation, and the '
  'appropriate electronic signatures.')

H('Overview')
P('The SMPL: Daily Observations XStep provides the operator with a Data Input table in which one row is '
  'recorded per daily sample. For each day the operator presses the Record button to stamp the sample Date and '
  'Time, then records the bioreactor conditions: Temperature, RPM (agitation), pH (with its probe) and External '
  'pH, DO% (with its probe), O2 Consumption, CPE %, Cell Count Live and Cell Count Dead. Each row is completed '
  'with a Performed-By signature, and the step is closed with a Witness signature.')
P('The sample Date and Time are populated by a system-time function module when the operator presses the Record '
  'button; these fields are read-only and are never hand-keyed. The numeric conditions that carry a specification '
  '(Temperature 25-29 degC, DO 20-100 %, RPM 55) are range-validated on entry; if a value is out of range the '
  'entry is only accepted once the Supervisor is contacted, a supervisory signature is captured and the operator '
  'comments on the corrective action taken. The pH and External pH entries are validated as real pH values '
  '(0.0-14.0) so an impossible pH cannot be recorded. The pH and DO probes are recorded as operator entries '
  '(they are not order-assigned equipment). Additional rows are added with the Add Row action, one per '
  'observation day.')
P('The bioreactor conditions are captured in the EBR Data Input table and are available for review and '
  'reporting. All function modules used by this XStep already exist in DE1 100 and are reused as-is. Note: the '
  'existing daily-sample persistence function module (ZSMPL_FM_STORE_DAILY_SAMPLE) stores the Seed / Production '
  'sampling parameters (glucose, BRX weight, VCD, viability) and does not cover the viral daily-observation '
  'parameters (CPE %, live/dead cell counts, O2 consumption); if these values must also be persisted to a custom '
  'table for cross-day retrieval, a dedicated persistence function module analogous to ZSMPL_FM_STORE_DAILY_SAMPLE '
  'would be required (the only potential new development).')

H('Reasons for developing')
P('The SMPL: Daily Observations XStep was developed to replace the daily bioreactor conditions table repeated '
  'throughout the paper batch record (Manufacturing Directions). Stamping the sample date and time with a '
  'system-time function module removes transcription error; range-validating the recorded parameters means an '
  'out-of-range result is only accepted with a supervisory signature and a corrective-action comment, making the '
  'paper "contact the Supervisor" instruction an enforced electronic control; and capturing electronic '
  'Performed-By and Witness signatures satisfies the data-integrity and dual-verification requirements. The '
  'single reusable block replaces the table that is otherwise reprinted for every observation day.')

H('Authorization')
P('Access to the XStep and the ability to enter data and apply signatures is controlled by the standard SiMPL '
  'EBR security model. The operator must hold the PFCG role assigned to the relevant process order / control '
  'recipe. Performed-By, Witness and Supervisor signatures are captured using the SAP digital signature method '
  'configured for the EBR.')

H('Assumptions/ Dependencies')
Bul('The generic SiMPL range, date/time, signature and activation function modules '
    '(/SMPL/PPPI_FM_MIN_MAX, /SMPL/PPPI_FM_GET_DATE_TIME, /SMPL/PPPI_FM_SIG_ADD_DB_CB, '
    '/SMPL/PPPI_FM_SIG_POPULATE_CB, /SMPL/PPPI_FM_SIG_VALIDATION, /SMPL/PPPI_FM_VALI_SUPE_SIG, '
    '/SMPL/PPPI_FM_INITIAL_ACTIVE, /SMPL/MBR_DEP_CHECK_ACTIVE) exist and are active in the target system.')
Bul('The pH and DO probe references are recorded as operator entries; the step does not read order-assigned '
    'equipment, so no equipment function module is required.')
Bul('The in-range specifications are configured on the step: Temperature 25-29 degC, DO 20-100 %, RPM 55 '
    '(setpoint), and pH 0.0-14.0 (validity). Out-of-range temperature/DO/RPM entries require a supervisory '
    'signature and a corrective-action comment.')
Bul('Persistence of the viral-specific observation parameters (CPE %, live/dead cell counts, O2 consumption) '
    'to a custom table, if required, is a separate development (see Overview).')

H('Validation Checks')
P('The following validations are applied within the XStep:')
T3(['Field','Validation','Function Module'],
   [['Date / Time','System-stamped (read-only) when the Record button is pressed','/SMPL/PPPI_FM_GET_DATE_TIME'],
    ['Temperature (degC)','Must be within 25-29 degC; out-of-range only accepted with a supervisor signature','/SMPL/PPPI_FM_MIN_MAX'],
    ['RPM','Must equal the configured setpoint (55); out-of-range only accepted with a supervisor signature','/SMPL/PPPI_FM_MIN_MAX'],
    ['DO %','Must be within 20-100 %; out-of-range only accepted with a supervisor signature','/SMPL/PPPI_FM_MIN_MAX'],
    ['pH / External pH','Must be a valid pH value (0.0 - 14.0)','/SMPL/PPPI_FM_MIN_MAX'],
    ['Performed By','Mandatory electronic signature per row','/SMPL/PPPI_FM_SIG_ADD_DB_CB / /SMPL/PPPI_FM_SIG_VALIDATION'],
    ['Witness By','Mandatory witness (supervisory) signature at step level','/SMPL/PPPI_FM_VALI_SUPE_SIG']])
BL()

H('XStep Layout Design')
P('The XStep is rendered as a Data Input table with an instruction panel. The instruction reads: "Sample the '
  'reactor daily and record bioreactor conditions in the Daily Observations table; stamp the sample date/time '
  'each day. Ensure parameters are in range (Temp 25-29 degC, DO 20-100 %, RPM 55). If out of range, contact '
  'the Supervisor and comment on any corrective action taken." Each table row captures the following:')
Bul('Record - button that stamps the sample Date and Time.')
Bul('Date / Time - read-only, system-stamped.')
Bul('Temp (degC) - entry, numeric, range-validated (25-29).')
Bul('RPM - entry, numeric, validated against the setpoint (55).')
Bul('pH / Probe - entry: measured pH and the pH probe reference, validated as a real pH value (0.0-14.0).')
Bul('External pH - entry: offline reference pH, validated as a real pH value (0.0-14.0).')
Bul('DO% / Probe - entry: measured dissolved oxygen and the DO probe reference, range-validated (20-100).')
Bul('O2 Consumption (L) - entry, numeric.')
Bul('CPE % - entry, numeric (cytopathic effect).')
Bul('Cell Count Live / Cell Count Dead - entry, numeric.')
Bul('Performed By - signature, required (one per row).')
P('Additional rows can be added with Add Row, one per observation day. A step-level Witness By signature is '
  'captured in the footer; out-of-range entries additionally require a Supervisor signature.')
add(IMG(OUTDIR+'/image.png'))
BL()

H('Function Module(s)')
P('The following function modules are used by this XStep. All of them already exist in DE1 100 and are reused '
  'as-is; no new function module is required for the reused behaviour.')
for fm in FM_ORDER:
    for el in B.lib_fm_elems(fm): add(el)

H('Pseudocode')
for fm in FM_ORDER:
    for el in B.lib_pseudo_elems(fm): add(el)

H('Configuration Specifications')
P('The XStep is configured using the following parameters:')
T3(['Parameter','Value / Configuration','Description'],
   [['Header','Daily Observations','Header text displayed at the top of the step (configurable in SiMPL MBR).'],
    ['Instruction','(instruction text)','The instruction text displayed to the operator (configurable in SiMPL MBR).'],
    ['Record button','Stamp Date/Time','Calls /SMPL/PPPI_FM_GET_DATE_TIME to populate the sample Date and Time.'],
    ['Temperature range','25 - 29 degC','Min/max limits enforced by /SMPL/PPPI_FM_MIN_MAX.'],
    ['DO range','20 - 100 %','Min/max limits enforced by /SMPL/PPPI_FM_MIN_MAX.'],
    ['RPM setpoint','55','Agitation setpoint validated by /SMPL/PPPI_FM_MIN_MAX.'],
    ['pH range','0.0 - 14.0','Valid pH range enforced by /SMPL/PPPI_FM_MIN_MAX (applies to pH and External pH).'],
    ['Performed By','Signature column','Per-row Performed-By signature (/SMPL/PPPI_FM_SIG_ADD_DB_CB).'],
    ['Witness By','Step-level signature','Witness signature (/SMPL/PPPI_FM_VALI_SUPE_SIG).'],
    ['Add Row','Enabled','One row per observation day.']])
BL()

H('Test Scenarios')
T3(['#','Scenario','Expected Result'],
   [['1','Press the Record button.','The sample Date and Time are stamped with the current system date and time (read-only).'],
    ['2','Enter Temperature = 27, RPM = 55, DO = 60 (all in range).','Values are accepted; no supervisor signature is required.'],
    ['3','Enter Temperature = 32 (out of range).','A range warning is raised (/SMPL/PPPI_FM_MIN_MAX); a supervisor signature and corrective-action comment are required.'],
    ['4','Enter a DO % below 20 or above 100.','The value is flagged as out of range; supervisor escalation is required.'],
    ['5','Enter a pH (or External pH) outside 0.0 - 14.0 (e.g. 15).','The value is rejected as not a valid pH (/SMPL/PPPI_FM_MIN_MAX, VALUE_OUT_OF_RANGE).'],
    ['6','Add a second row with Add Row.','A new observation-day row is added and can be completed independently.'],
    ['7','Attempt to complete the step without a Performed-By or Witness signature.','The step cannot be completed; both signatures are mandatory.'],
    ['8','Re-open the completed step.','The recorded observations, stamped date/time and signatures are displayed and retained.']])
BL()

H('Document References')
T2(['Reference No.','Document Title'],
   [['1','Manufacturing Directions - Sample Virus in 2000L Single-Use Bioreactor (SUB)'],
    ['2','SiMPL XStep Library - Process Manufacturing (DE1 100)'],
    ['3','SMPL: Bioreactor Sampling XStep Design Specification (reference archetype)']])
BL()

H('Revision History')
T3(['Version No.','Description','Date'],
   [['1.0','Initial document','2026-07-07']])

# ---------------- fixes + save ----------------
B.neutralise_empty_headings(body)
B.set_update_fields(doc)
doc.save(OUT)

# title-canvas + header/footer text: rename the XStep
tmp=OUT+'.tmp'
with zipfile.ZipFile(OUT) as z: data={n:z.read(n) for n in z.namelist()}
for n in list(data):
    if re.match(r'word/(document|header\d+|footer\d+)\.xml$', n):
        x=data[n].decode('utf-8')
        x2=x.replace('Three Variable Calc','Daily Observations')
        if x2!=x: data[n]=x2.encode('utf-8')
with zipfile.ZipFile(tmp,'w',zipfile.ZIP_DEFLATED) as z:
    for n,b in data.items(): z.writestr(n,b)
shutil.move(tmp,OUT)
sys.stdout.flush()
print('built',OUT)
