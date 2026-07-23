# -*- coding: utf-8 -*-
"""Extract the FULL text of the OLD paper batch records (the two .docx source MPRs) so the
'New vs Old' EBR can show the real instruction text beneath each XStep — a true fulfillment
comparison, not the workbook's one-line summary.

Walks the document body in order, tracks the current Heading-1 section (§1..§n as they appear),
and within each section groups the content the way the paper record reads:
  - an INSTRUCTION (an imperative action sentence), followed by
  - the FIELDS that instruction makes the operator record (pH, weights, IDs, times, ...).
Word auto-numbers steps and python-docx cannot read those list numbers, so we DO NOT fabricate
step numbers; we keep the real section number (§7..§n) and present the verbatim content beneath it.

Exposes: SECTIONS[tag][sec_no] = {'title': str, 'groups': [ {'instr': str, 'fields': [str,...]}, ... ]}
         tag = 'VI' | 'CDF'
"""
import docx, re
from docx.oxml.ns import qn

BRDIR = r"c:\Users\carlo\Dev\TechSpecs\AZ Phase 3 Virus Inactivation and filtration\Batch Record"
DOCS = {
 'VI':  BRDIR + r"\PN- 8012441 - AZD0543 -  Virus Inactivation (Low pH, GPF, 2000 L Scale, Process 1).docx",
 'CDF': BRDIR + r"\PN- 8012475 - AZD0543 - Concentration and Diafiltration (Large Skid, GPF, 2000 L Scale, Process 1).docx",
}

# signature/header/boilerplate cells to drop outright
_SKIP = re.compile(r'^(step|description|mpr step no\.?|part no\.?|batch no\.?|exp\.? date|'
                   r'performed\s*/?\s*by\s*/?\s*date|witnessed\s*/?\s*by\s*/?\s*date|'
                   r'performed by\s*/?\s*date|witnessed by\s*/?\s*date|sap consumption.*|'
                   r'material description|autoclave.*|comments?|initials?|verified by\s*/?\s*date|'
                   r'value|result|units?|uom|n/?a|verified by|checked by|parameter|valve id|'
                   r'setpoint|alarm|valve state|uf/?df system|opened|closed)$', re.I)

# an instruction starts with (or clearly is) an imperative / conditional action sentence.
# \b matters: "Mixer"/"Mixing" must NOT match "mix", "Weight" must NOT match "weigh", etc.
_VERB = re.compile(r'^(as needed|obtain|record|attach|confirm|ensure|calculate|determine|begin|'
                   r'remove|aliquot|submit|measure|store|mix|use|verify|open|close|charge|'
                   r'recirculate|disconnect|drain|keep|prepare|choose|document|indicate|'
                   r'transfer|note|on the|to prepare|to prep|prior to|slowly|follow|if|'
                   r'for|when|after|once|select|complete|perform|input|click|download|run|'
                   r'stop|collect|label|place|repeat|proceed|do not|add|take|weigh|'
                   r'the following|steps? \d)\b', re.I)

def _clean(s):
    return re.sub(r'\s+', ' ', (s or '').replace('\xa0', ' ')).strip()

def _dedupe_merged(cells):
    out = []
    for c in cells:
        if not out or out[-1] != c:
            out.append(c)
    return out

def _is_noise_field(t):
    """A field line carrying no real label — only numbers/UoM/blanks/ranges."""
    core = re.sub(r'[\d\.\,\;\:\(\)\/\=\-\_°º\s]', '', t)
    core = re.sub(r'\b(kg|L|mL|g|hrs?|mins?|min|mS|cm|oC|C|RPM|psig|target|max|min|range|to)\b',
                  '', core, flags=re.I)
    return len(core) < 3

def _is_instruction(t):
    if _VERB.match(t):
        return True
    # long sentence-like line with spaces and a trailing period
    return len(t) > 70 and ' ' in t and t.rstrip().endswith('.')

def _row_lines(row):
    cells = _dedupe_merged([_clean(c.text) for c in row.cells])
    return [c for c in cells if c and not _SKIP.match(c) and len(c) > 2]

def _iter_body(doc):
    body = doc.element.body
    tbls = iter(doc.tables)
    paras = {p._p: p for p in doc.paragraphs}
    for child in body.iterchildren():
        if child.tag == qn('w:p'):
            p = paras.get(child)
            if p is None:
                continue
            txt = _clean(p.text)
            if not txt:
                continue
            yield ('h1', txt) if p.style.name == 'Heading 1' else ('p', txt)
        elif child.tag == qn('w:tbl'):
            yield ('tbl', next(tbls))

def extract(tag):
    doc = docx.Document(DOCS[tag])
    sections, sec_no, cur = {}, 0, None

    def flush_line(txt, force_instr=False):
        if txt.strip().lower() in ('as needed', 'as required', 'n/a', 'as applicable'):
            return                                       # BOM quantity annotations, not steps
        if _is_instruction(txt) or force_instr:
            cur['groups'].append({'instr': txt, 'fields': []})
        else:
            if _is_noise_field(txt):
                return
            if not cur['groups']:
                cur['groups'].append({'instr': '(records)', 'fields': []})
            cur['groups'][-1]['fields'].append(txt)

    for item in _iter_body(doc):
        if item[0] == 'h1':
            sec_no += 1
            cur = {'title': item[1], 'groups': []}
            sections[sec_no] = cur
        elif cur is None:
            continue
        elif item[0] == 'p':
            txt = item[1]
            if len(txt) > 25 and not txt.lower().startswith('table continued'):
                flush_line(txt, force_instr=True)  # narrative notes read as instructions
        elif item[0] == 'tbl':
            for row in item[1].rows:
                for ln in _row_lines(row):
                    flush_line(ln)
    return sections

SECTIONS = {tag: extract(tag) for tag in DOCS}

# ---- Which XStep card each OLD process section / attachment belongs under ----
# Sections 1-6 (summary/flowchart/refs/BOM/notes) and the revision history are context, not
# recordable process steps, so they are intentionally omitted. Reusable "Common" blocks that are
# instantiated *inside* several sections (Equipment ID, Measure pH, Weigh, Calculation, Material
# Addition, Product Transfer) have no single home section and fall back to the workbook summary.
SEC2FOLDER = {
 'VI': {
   4:  "Common - Display BOM",
   5:  "Common - Additional Solution Batches",
   6:  "Common - Process Notes and Global Limits",
   7:  "Common - Incoming Product Information",
   8:  "VI - Treatment Vessel Setup",
   9:  "VI - Temperature Decision Tree",
   10: "VI - Acid-Base pH Titration Table",
   11: "VI - Incubation Timing and Sample",
   12: "VI - Acid-Base pH Titration Table",
   13: "VI - Incubation Timing and Sample",
   14: "VI - Acid-Base pH Titration Table",
   15: "Common - Product Sampling and DLIMS Submission",
   16: "Common - Instruction and Sign-off",
   17: "Common - Product Sampling and DLIMS Submission",
   18: "Common - Non-Routine Sampling Record",
   19: "Common - Hold Time Table",
   20: "Common - Yield Calculations",
   21: "Common - Transfer Tubing and Hosing Info",
   22: "Common - Product Recirculation Worksheet",
   23: "Common - Comments and Deviations",
 },
 'CDF': {
   4:  "Common - Display BOM",
   5:  "Common - Additional Solution Batches",
   6:  "Common - Process Notes and Global Limits",
   7:  "CDF - Minimum Membrane Surface Area",
   8:  "CDF - Run Skid Recipe",
   9:  "Common - Hold Time Table",
   10: "CDF - Run Skid Recipe",
   11: "CDF - Record CIPDS Skid Cassette Info",
   12: "CDF - Process Input Calculations",
   13: "CDF - Run Skid Recipe",
   14: "CDF - Continued Diafiltration Decision",
   15: "CDF - Run Skid Recipe",
   16: "CDF - Recovery Calculations and Execution",
   17: "CDF - Dilution Decision and Calculations",
   18: "CDF - PFI Storage Vessel and Filtration",
   19: "CDF - Post-Processing and Sanitization",
   20: "Common - Instruction and Sign-off",
   21: "Common - Transfer Tubing and Hosing Info",
   22: "CDF - Operation Monitoring Worksheet",
   23: "Common - Product Sampling and DLIMS Submission",
   24: "Common - Non-Routine Sampling Record",
   25: "Common - Hold Time Table",
   26: "CDF - Record CIPDS Skid Cassette Info",
   27: "Common - Product Recirculation Worksheet",
   28: "Common - Yield Calculations",
   29: "Common - Comments and Deviations",
 },
}

def full_for_folder(folder):
    """[(tag, sec_no, section_title, groups), ...] of OLD sections mapped to this XStep, in order."""
    out = []
    for tag in ('VI', 'CDF'):
        for sec_no, fld in sorted(SEC2FOLDER[tag].items()):
            if fld == folder:
                sec = SECTIONS[tag].get(sec_no)
                if sec and sec['groups']:
                    out.append((tag, sec_no, sec['title'], sec['groups']))
    return out

# ---- Cross-cutting reusable blocks: no single home section, so harvest the verbatim instruction
# groups from wherever they occur (matched on the instruction text + its captured fields). ----
# NOTE: matched against the INSTRUCTION text only (not the captured fields) — a group belongs to a
# reusable block only when the ACTION is about that activity, not when it merely records e.g. a
# scale ID in passing (that field still shows, in context, under its own action).
HARVEST_KW = {
 "Common - Room and Equipment Assign": [
   "record the room", "confirm that the ph meter", "confirm that the thermometer",
   "confirm that the conductivity", "re-standardiz", "restandardiz", "standardization of",
   "recalibrate the ph meter", "meter standardization", "verify skid meter",
   "confirm the ph meter", "confirm the thermometer", "record the equipment id",
   "within its calibration"],
 "Common - Product pH Conductivity Temperature": [
   "measure the ph", "measure the conductivity", "measure the ph, conductivity",
   "ph, conductivity, and temperature", "measure the final ph", "record the final ph",
   "measure the offline ph"],
 "Common - Product Vessel Weigh": [
   "obtain the gross", "obtain the tare", "obtain the gross and net", "record the gross weight",
   "record the tare weight", "record the net weight", "weigh the", "obtain the net weight"],
 "Common - Product Mixing": [
   "mix the product", "ensure agitation", "additional mixing", "mix the incubation",
   "mix for at least"],
 "Common - Calc Three Columns": [
   "calculate the", "determine the initial", "calculate the expected"],
 "Common - Additional Manufacturing Supplies": [
   "ensure the material is acceptable", "additional material", "record the applicable material"],
 "Common - Incoming Product Information": [
   "record the tare weight and net weight", "solovpe concentration", "record vf product information",
   "record the following information from the azd0543", "virus filtration product information",
   "affinity product solovpe", "product information"],
}

def _norm(t):
    return re.sub(r'\s+', ' ', t.lower()).strip()

def harvest_for_folder(folder):
    """Verbatim groups matching this reusable block's keywords, deduped, tagged with section.
    Returns [(tag, sec_no, section_title, [group,...]), ...] grouped by section, in doc order."""
    kws = HARVEST_KW.get(folder)
    if not kws:
        return []
    seen, by_sec = set(), {}
    for tag in ('VI', 'CDF'):
        for sec_no in sorted(SECTIONS[tag]):
            if sec_no < 7:                      # §1-6 = summary/flowchart/refs/BOM/notes (context)
                continue
            sec = SECTIONS[tag][sec_no]
            if 'REVISION HISTORY' in sec['title'].upper():
                continue
            for g in sec['groups']:
                if len(g['instr']) < 6 or g['instr'] == '(records)':
                    continue
                hay = _norm(g['instr'])                 # match on the ACTION only, not fields
                if any(k in hay for k in kws):
                    key = _norm(g['instr'])[:80]
                    if key in seen:
                        continue
                    seen.add(key)
                    by_sec.setdefault((tag, sec_no, sec['title']), []).append(g)
    return [(t, n, ttl, gs) for (t, n, ttl), gs in by_sec.items()]

def _count(tag):
    n_i = n_f = 0
    for sec in SECTIONS[tag].values():
        for g in sec['groups']:
            n_i += 1; n_f += len(g['fields'])
    return n_i, n_f

if __name__ == '__main__':
    for tag, secs in [('VI', [10, 15, 19]), ('CDF', [12, 16])]:
        for s in secs:
            sec = SECTIONS[tag].get(s)
            print(f'\n===== {tag} §{s}  {sec["title"]}  ({len(sec["groups"])} instr groups) =====')
            for g in sec['groups'][:14]:
                print('  * ' + g['instr'][:110])
                for f in g['fields'][:8]:
                    print('       - ' + f[:90])
    for tag in DOCS:
        i, f = _count(tag)
        print(f'\n{tag}: {i} instruction-groups, {f} field lines across {len(SECTIONS[tag])} sections')
