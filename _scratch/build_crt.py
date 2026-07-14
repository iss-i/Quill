# -*- coding: utf-8 -*-
"""Build the SMPL: Cell Receipt & Transfer XStep Design Specification (.docx) in the
AZ Phase 2 template. Clones the Three Variable Calc doc for the template shell (title
page, styles, live TOC), rebuilds the body with Cell Receipt & Transfer content, and
reuses the Phase-2 FM/pseudocode emitters. All 8 FMs already exist in DE1 100; only
ZSMPL_FM_CHECK_CHAR_DATE is authored here (sourced from DE1 100) and injected into the
FM library at runtime."""
import sys, io, copy, os, shutil, zipfile, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import build_phase2 as B
from p2lib import is_heading, heading_text

ROOT='..'
DONOR=ROOT+'/Phase 2 XSteps/SMPL- Three Variable Calc/SMPL_Three Variable Calc XStep Design Specification.docx'
OUTDIR=ROOT+'/Merck XStep Mock Ups/Cell Receipt and Transfer'
OUT=OUTDIR+'/SMPL_Cell Receipt & Transfer XStep Design Specification.docx'

# ---- author ZSMPL_FM_CHECK_CHAR_DATE from DE1 100 source, inject into library ----
B.FMDATA['ZSMPL_FM_CHECK_CHAR_DATE']={
 'name':'ZSMPL_FM_CHECK_CHAR_DATE',
 'imports':[
   ('IV_VALUE','CHAR10','Date value entered by the operator, in yyyy.mm.dd format (or "N/A") (optional)'),
   ('IV_DATE','DATS','Date value supplied directly in internal (DATS) format (optional)'),
 ],
 'exceptions':[
   ('INCORRECT_DATE_FORMAT','Entered value is not in the expected yyyy.mm.dd format'),
   ('INCORRECT_VALUE_INPUT','Entered value is neither a valid date nor "N/A"'),
   ('DATE_CANNOT_BE_IN_PAST','The entered/derived date is earlier than the current system date'),
 ],
 'pseudo':[
   ('h','1. Purpose'),
   ('p','Validates a character-format date entered by the operator (for example the Sterilization Exp. Date). '
        'Confirms the value is either a valid date in yyyy.mm.dd format or the literal "N/A", and that any date '
        'supplied is not earlier than the current system date.'),
   ('h','2. Input Parameters'),
   ('b','IV_VALUE - the date the operator typed, as a 10-character string in yyyy.mm.dd format, or "N/A" / "NA" (optional).'),
   ('b','IV_DATE - an alternative date supplied directly in internal (DATS) format (optional).'),
   ('h','3. Processing Logic'),
   ('s','a. Take a working copy of IV_VALUE, translate it to upper case and measure its length; build an internal '
        'date value by removing the "." separators.'),
   ('s','b. If IV_VALUE is not blank:'),
   ('b','If it equals "N/A" or "NA", accept the entry and perform no further checks.'),
   ('b','Otherwise, if the month portion (positions 6-7) is greater than 12, raise INCORRECT_DATE_FORMAT with the '
        'message "Incorrect date format entered - use yyyy.mm.dd".'),
   ('b','Otherwise, if the value is shorter than 10 characters, raise INCORRECT_VALUE_INPUT with the message '
        '"Incorrect value input - options are Date or N/A".'),
   ('b','Otherwise, if the resulting date is earlier than the system date, raise DATE_CANNOT_BE_IN_PAST.'),
   ('s','c. If IV_DATE is supplied and is earlier than the system date, raise DATE_CANNOT_BE_IN_PAST.'),
   ('h','Additional Notes'),
   ('h','Error Handling:'),
   ('b','INCORRECT_DATE_FORMAT - the month component is invalid / the value is not in yyyy.mm.dd format.'),
   ('b','INCORRECT_VALUE_INPUT - the value is neither a full date nor "N/A".'),
   ('b','DATE_CANNOT_BE_IN_PAST - the entered or derived date precedes the current system date.'),
   ('h','Dependencies:'),
   ('b','None - pure ABAP date handling; no database reads or external function-module calls.'),
   ('h','Assumptions:'),
   ('b','Operators enter dates in yyyy.mm.dd format; "N/A" is permitted where a sterilization expiry does not apply.'),
 ],
}

FM_ORDER=['ZSMPL_FM_CUSTOM_INDEX','/SMPL/PPPI_FM_GET_DATE_TIME','ZSMPL_FM_CHECK_CHAR_DATE',
          '/SMPL/PPPI_FM_SIG_ADD_DB_CB','/SMPL/PPPI_FM_VALI_SUPE_SIG',
          '/SMPL/PPPI_FM_INITIAL_ACTIVE','/SMPL/MBR_DEP_ADD_PERFORM','/SMPL/MBR_DEP_CHECK_ACTIVE']

os.makedirs(OUTDIR, exist_ok=True)
shutil.copy(DONOR, OUT)
doc=Document(OUT); body=doc.element.body
children=list(body)
sectPr=children[-1] if children[-1].tag==qn('w:sectPr') else None
purpose=next(el for el in children if el.tag==qn('w:p') and is_heading(el) and heading_text(el).strip()=='Purpose')
anchor=copy.deepcopy(purpose)
start=children.index(purpose); end=children.index(sectPr) if sectPr is not None else len(children)
for el in children[start:end]: body.remove(el)

bid=B.max_bookmark_id(body)+1
def add(el):
    if sectPr is not None: sectPr.addprevious(el)
    else: body.append(el)
def H(t):
    global bid
    add(B.make_heading_like(anchor,t,bid)); bid+=1
def P(t): add(B.mk_plain(t))
def Bul(t): add(B.mk_bullet(t))
def BL(): add(B.mk_blank())
def TBL(tmpl, headers, rows):
    t=B.mk_table(tmpl, rows); hdr=t.findall(qn('w:tr'))[0]; tcs=hdr.findall(qn('w:tc'))
    for i,h in enumerate(headers): B.set_cell_text(tcs[i], h)
    add(t)
def T3(headers, rows): TBL(B.T_PTABLE, headers, rows)
def T2(headers, rows): TBL(B.T_ETABLE, headers, rows)
def IMG(path, width_in=6.5):
    para=doc.add_paragraph(); para.alignment=WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(path, width=Inches(width_in))
    el=para._p; el.getparent().remove(el); return el

# ---------------- content ----------------
H('Purpose')
P('The purpose of this document is to outline the design and configuration of the SMPL: Cell Receipt & Transfer '
  'XStep. This XStep records the receipt of cells from the cell prep lab and their transfer into the production '
  'bioreactor, capturing the receipt date and time, the cell tank number and its sterilization expiry date, and '
  'the transfer time, with the appropriate electronic signatures.')

H('Overview')
P('The SMPL: Cell Receipt & Transfer XStep provides the operator with a data-input table that documents the '
  'receipt and transfer of cells. When the cells are received, the operator presses the Record button to stamp '
  'the Date Received and Time Received; records the Cell Tank number and the vessel Sterilization Exp. Date; and, '
  'once the media has been confirmed at the required temperature and the cells transferred, presses a second '
  'Record button to stamp the Transfer Time. Each row is completed with a Performed-By signature, and the step '
  'is closed with a Witness signature.')
P('The Date Received, Time Received and Transfer Time are populated by a system-time function module when the '
  'operator presses the associated Record button; these fields are read-only and are never hand-keyed. The '
  'Sterilization Exp. Date is validated as a date. Additional rows can be added with the Add Row action, and the '
  'rows are numbered automatically.')
P('All function modules used by this XStep already exist in DE1 100; no new development is required.')

H('Reasons for developing')
P('The SMPL: Cell Receipt & Transfer XStep was developed to replace the manual receipt-and-transfer record in '
  'the paper batch record (Manufacturing Directions). Stamping the receipt and transfer times with a system-time '
  'function module removes transcription error and guarantees an accurate, tamper-evident timestamp; validating '
  'the sterilization expiry date and capturing electronic Performed-By and Witness signatures enforces the '
  'data-integrity and dual-verification requirements of the process.')

H('Authorization')
P('Access to the XStep and the ability to enter data and apply signatures is controlled by the standard SiMPL '
  'EBR security model. The operator must hold the PFCG role assigned to the relevant process order / control '
  'recipe. Performed-By and Witness signatures are captured using the SAP digital signature method configured '
  'for the EBR.')

H('Assumptions/ Dependencies')
Bul('The generic SiMPL date/time, indexing, signature and activation function modules '
    '(/SMPL/PPPI_FM_GET_DATE_TIME, ZSMPL_FM_CUSTOM_INDEX, ZSMPL_FM_CHECK_CHAR_DATE, '
    '/SMPL/PPPI_FM_SIG_ADD_DB_CB, /SMPL/PPPI_FM_VALI_SUPE_SIG, /SMPL/PPPI_FM_INITIAL_ACTIVE, '
    '/SMPL/MBR_DEP_ADD_PERFORM, /SMPL/MBR_DEP_CHECK_ACTIVE) exist and are active in the target system.')
Bul('The cell prep information is available from the cell prep lab and is attached to the Manufacturing '
    'Directions as a separate step.')
Bul('The media has been confirmed at the required temperature (no warmer than the standard 2-8 degC) before '
    'the cells are transferred; up to the standard 10 lb air pressure may be applied.')
Bul('The Sterilization Exp. Date is recorded in yyyy.mm.dd format, or "N/A" where a sterilization expiry does '
    'not apply.')

H('Validation Checks')
P('The following validations are applied within the XStep:')
T3(['Field','Validation','Function Module'],
   [['Date Received','System-stamped (read-only) when the Record button is pressed','/SMPL/PPPI_FM_GET_DATE_TIME'],
    ['Time Received','System-stamped (read-only) when the Record button is pressed','/SMPL/PPPI_FM_GET_DATE_TIME'],
    ['Cell Tank #','Free-text entry; required','(none)'],
    ['Sterilization Exp. Date','Must be a valid date in yyyy.mm.dd format (or "N/A") and not in the past','ZSMPL_FM_CHECK_CHAR_DATE'],
    ['Transfer Time','System-stamped (read-only) when the Record button is pressed','/SMPL/PPPI_FM_GET_DATE_TIME'],
    ['Performed By','Mandatory electronic signature per row','/SMPL/PPPI_FM_SIG_ADD_DB_CB'],
    ['Witness By','Mandatory witness (supervisory) signature at step level','/SMPL/PPPI_FM_VALI_SUPE_SIG']])
BL()

H('XStep Layout Design')
P('The XStep is rendered as a Data Input table with an instruction panel. The instruction reads: "Receive cells '
  'from the cell prep lab and attach the cell prep information to this MD. Ensure the media is at the required '
  'temperature before transferring cells (up to 10 lb air pressure may be applied). Stamp the receipt and transfer '
  'times; record cell tank and sterilization expiry." Each table row captures the following:')
Bul('Record (receipt) - button that stamps Date Received and Time Received.')
Bul('Date Received - read-only, system-stamped.')
Bul('Time Received - read-only, system-stamped.')
Bul('Cell Tank # - entry, required.')
Bul('Sterilization Exp. Date - entry, validated as a date (yyyy.mm.dd or "N/A").')
Bul('Record (transfer) - button that stamps Transfer Time.')
Bul('Transfer Time - read-only, system-stamped.')
Bul('Performed By - signature, required (one per row).')
P('Additional rows can be added with Add Row; the rows are numbered automatically. A step-level Witness By '
  'signature is captured in the footer.')
add(IMG(OUTDIR+'/image.png'))
BL()

H('Function Module(s)')
P('The following function modules are used by this XStep. All of them already exist in DE1 100 and are reused '
  'as-is; no new function module is required.')
for fm in FM_ORDER:
    for el in B.lib_fm_elems(fm): add(el)

H('Pseudocode')
for fm in FM_ORDER:
    for el in B.lib_pseudo_elems(fm): add(el)

H('Configuration Specifications')
P('The XStep is configured using the following parameters:')
T3(['Parameter','Value / Configuration','Description'],
   [['Header','Cell Receipt & Transfer','Header text displayed at the top of the step (configurable in SiMPL MBR).'],
    ['Instruction','(instruction text)','The instruction text displayed to the operator (configurable in SiMPL MBR).'],
    ['Record (receipt) button','Stamp Date/Time Received','Calls /SMPL/PPPI_FM_GET_DATE_TIME to populate Date Received and Time Received.'],
    ['Cell Tank #','Entry field','Free-text entry for the cell tank identifier (required).'],
    ['Sterilization Exp. Date','Entry field, date validator','Validated by ZSMPL_FM_CHECK_CHAR_DATE (yyyy.mm.dd or "N/A", not in past).'],
    ['Record (transfer) button','Stamp Transfer Time','Calls /SMPL/PPPI_FM_GET_DATE_TIME to populate Transfer Time.'],
    ['Performed By','Signature column','Per-row Performed-By signature (/SMPL/PPPI_FM_SIG_ADD_DB_CB).'],
    ['Witness By','Step-level signature','Witness/supervisory signature (/SMPL/PPPI_FM_VALI_SUPE_SIG).'],
    ['Add Row','Enabled','Rows numbered automatically by ZSMPL_FM_CUSTOM_INDEX.']])
BL()

H('Test Scenarios')
T3(['#','Scenario','Expected Result'],
   [['1','Press the receipt Record button.','Date Received and Time Received are stamped with the current system date and time (read-only).'],
    ['2','Enter the Cell Tank number and a valid Sterilization Exp. Date (e.g. 2027.01.15).','Values are accepted; the date passes validation.'],
    ['3','Enter a Sterilization Exp. Date with an invalid month (e.g. 2027.13.01).','Entry is rejected with "Incorrect date format entered - use yyyy.mm.dd" (INCORRECT_DATE_FORMAT).'],
    ['4','Enter a Sterilization Exp. Date that is in the past.','Entry is rejected with "Date cannot be in the past" (DATE_CANNOT_BE_IN_PAST).'],
    ['5','Enter "N/A" for the Sterilization Exp. Date.','The entry is accepted (no expiry applies).'],
    ['6','Press the transfer Record button.','Transfer Time is stamped with the current system time (read-only).'],
    ['7','Add a second row with Add Row.','A new numbered row is added and can be completed independently.'],
    ['8','Attempt to complete the step without a Performed-By or Witness signature.','The step cannot be completed; both signatures are mandatory.'],
    ['9','Re-open the completed step.','The recorded values, stamped times and signatures are displayed and retained.']])
BL()

H('Document References')
T2(['Reference No.','Document Title'],
   [['1','Manufacturing Directions - Sample Virus in 2000L Single-Use Bioreactor (SUB)'],
    ['2','SiMPL XStep Library - Process Manufacturing (DE1 100)']])
BL()

H('Revision History')
T3(['Version No.','Description','Date'],
   [['1.0','Initial document','2026-07-07']])

# ---------------- fixes + save ----------------
B.neutralise_empty_headings(body)
B.set_update_fields(doc)
doc.save(OUT)

# title-canvas + header/footer text: rename the XStep (XML-escape the ampersand)
tmp=OUT+'.tmp'
with zipfile.ZipFile(OUT) as z: data={n:z.read(n) for n in z.namelist()}
for n in list(data):
    if re.match(r'word/(document|header\d+|footer\d+)\.xml$', n):
        x=data[n].decode('utf-8')
        x2=x.replace('Three Variable Calc','Cell Receipt &amp; Transfer')
        if x2!=x: data[n]=x2.encode('utf-8')
with zipfile.ZipFile(tmp,'w',zipfile.ZIP_DEFLATED) as z:
    for n,b in data.items(): z.writestr(n,b)
shutil.move(tmp,OUT)
print('built',OUT)
