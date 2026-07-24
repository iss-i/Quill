# -*- coding: utf-8 -*-
"""Build the SMPL: VI Treatment Vessel Setup XStep Design Specification (.docx) in the
AZ template. From-scratch build (AZ Phase 3 VI+CDF mock-up only): clones the Three
Variable Calc doc for the shell (title page, styles, live TOC), rebuilds the 13 sections
from the VI Treatment Vessel Setup mock-up, and reuses the Phase-2 FM/pseudocode emitters.

This is a bespoke composite XStep: Yes/No dropdowns (Use-Affinity-vessel? / Bag Required? /
Tank Required? / Vent Filter Required?) DRIVE conditional activation of the section beneath
each (PI-PCS requires the gate and the gated section to live in the SAME XStep), so it is
built as one New XStep rather than a stack of reuse references. Bag / vent-filter /
product-filter lines validate Part No.->Material Description and Batch No.->Exp. Date, record
autoclave info, raise a Z_PICONS goods issue on signing and carry a per-row Performed By; the
Tank section records Tank ID / CIP batch+expiry / a Get-Pressure-Date-Time stamp / SIP batch;
a Get Scale/Balance retrieves the balance ID + calibration due; the tare weight is recorded;
three Yes/No dropdowns confirm which filters were attached at tare; a Recorded By / Date
signature closes the step.

FM set (14, all in DE1 100 - no new development):
  ZSMPL_FM_CUSTOM_INDEX           product-filter # row index                 (FMDATA)
  ZSMPL_FM_GET_MAT_DETAIL         Part No -> Material Description (validation) (FMDATA)
  ZSMPL_FM_GET_EXPIRY_DATE        Batch No -> Exp. Date (validation)          (FMDATA)
  ZSMPL_FM_CHECK_CHAR_DATE        Autoclave Exp. Date validator               (injected)
  ZSMPL_FM_CREATE_PROC_MSG        goods issue Z_PICONS mvt 261 on signing     (FMDATA)
  ZSMPL_FM_GET_ASSIGNED_EQUI_EBR  Get Scale / Balance (assigned equipment)    (biosamp)
  /SMPL/PPPI_FM_GET_DATE_TIME     Get Pressure Date/Time stamp                (biosamp)
  /SMPL/PPPI_FM_SIG_ADD_DB_CB     Performed-By signature (per row)            (golden)
  /SMPL/PPPI_FM_SIG_POPULATE_CB   signature callback                         (golden)
  /SMPL/PPPI_FM_SIG_VALIDATION    signature validator                        (golden)
  /SMPL/PPPI_FM_VALI_SUPE_SIG     Recorded-By (verification) signature        (golden)
  /SMPL/PPPI_FM_INITIAL_ACTIVE    step + section conditional activation       (golden)
  /SMPL/MBR_DEP_ADD_PERFORM       Performed-By dependency                     (golden)
  /SMPL/MBR_DEP_CHECK_ACTIVE      active-check dependency                     (golden)

The Yes/No gate / attachment / filters-attached dropdowns are restricted-value CT04
characteristics (ZSMPL_CHAR_YES_NO = Yes / No) captured to a local value (LV_DROP) - no
function module. Conditional activation is done with setup Commands, not an FM: each gated
section carries TABLE.ACTIVATE / TABLE.DEACTIVATE (and PROC_INSTR.ACTIVATE / DEACTIVATE,
LOCK / UNLOCK) whose Formula triggers evaluate LV_DROP - DEACTIVATE on (LV_ACTIVE <> 1 AND
IV_ACT_FLP <> 1) OR (LV_ACTIVE = 1 AND IV_ACT_FLP = 1) OR (LV_DROP = 2), ACTIVATE on the
complementary formula (the SMPL: EFS Additional Staging pattern). INITIAL_ACTIVE governs the
step's base initial-active state.
"""
import sys, io, copy, os, shutil, zipfile, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import build_phase2 as B
from p2lib import is_heading, heading_text

ROOT='..'
DONOR=ROOT+'/Phase 2 XSteps/SMPL- Three Variable Calc/SMPL_Three Variable Calc XStep Design Specification.docx'
OUTDIR=ROOT+'/AZ Phase 3 Virus Inactivation and filtration/Functional Specs'
OUT=OUTDIR+'/SMPL_VI Treatment Vessel Setup XStep Design Specification.docx'
IMG_SRC=ROOT+'/AZ Phase 3 Virus Inactivation and filtration/XStep mockups/New XSteps/VI - Treatment Vessel Setup/image.png'

# ---------------------------------------------------------------------------
# CHECK_CHAR_DATE is not in the persistent FM library - inject it (sourced from
# DE1 100) so the emitters render it in the accepted style. The material / expiry /
# goods-issue / index FMs already live in fmdata.py; equipment, date-time and the
# signature/activation FMs resolve from the BIOSAMP / GOLDEN reference specs.
# ---------------------------------------------------------------------------
B.FMDATA['ZSMPL_FM_CHECK_CHAR_DATE']={
 'name':'ZSMPL_FM_CHECK_CHAR_DATE',
 'imports':[
   ('IV_VALUE','CHAR10','Date value entered by the operator, in yyyy.mm.dd format (or "N/A") (optional)'),
   ('IV_DATE','DATS','Date value supplied directly in internal (DATS) format (optional)'),
 ],
 'exceptions':[
   ('INCORRECT_DATE_FORMAT','Entered value is not in the expected yyyy.mm.dd format'),
   ('INCORRECT_VALUE_INPUT','Entered value is neither a valid date nor "N/A"'),
   ('DATE_CANNOT_BE_IN_PAST','The entered/derived date is earlier than the current system date'),
 ],
 'pseudo':[
   ('h','1. Purpose'),
   ('p','Validates a character-format date entered by the operator (for example the Autoclave Exp. Date). '
        'Confirms the value is either a valid date in yyyy.mm.dd format or the literal "N/A", and that any date '
        'supplied is not earlier than the current system date.'),
   ('h','2. Input Parameters'),
   ('b','IV_VALUE - the date the operator typed, as a 10-character string in yyyy.mm.dd format, or "N/A" / "NA" (optional).'),
   ('b','IV_DATE - an alternative date supplied directly in internal (DATS) format (optional).'),
   ('h','3. Processing Logic'),
   ('s','a. Take a working copy of IV_VALUE, translate it to upper case and measure its length; build an internal '
        'date value by removing the "." separators.'),
   ('s','b. If IV_VALUE is not blank:'),
   ('b','If it equals "N/A" or "NA", accept the entry and perform no further checks.'),
   ('b','Otherwise, if the month portion (positions 6-7) is greater than 12, raise INCORRECT_DATE_FORMAT with the '
        'message "Incorrect date format entered - use yyyy.mm.dd".'),
   ('b','Otherwise, if the value is shorter than 10 characters, raise INCORRECT_VALUE_INPUT with the message '
        '"Incorrect value input - options are Date or N/A".'),
   ('b','Otherwise, if the resulting date is earlier than the system date, raise DATE_CANNOT_BE_IN_PAST.'),
   ('s','c. If IV_DATE is supplied and is earlier than the system date, raise DATE_CANNOT_BE_IN_PAST.'),
   ('h','Additional Notes'),
   ('h','Error Handling:'),
   ('b','INCORRECT_DATE_FORMAT - the month component is invalid / the value is not in yyyy.mm.dd format.'),
   ('b','INCORRECT_VALUE_INPUT - the value is neither a full date nor "N/A".'),
   ('b','DATE_CANNOT_BE_IN_PAST - the entered or derived date precedes the current system date.'),
   ('h','Dependencies:'),
   ('b','None - pure ABAP date handling; no database reads or external function-module calls.'),
   ('h','Assumptions:'),
   ('b','Operators enter dates in yyyy.mm.dd format; "N/A" is permitted where an autoclave expiry does not apply.'),
 ],
}

FM_ORDER=['ZSMPL_FM_CUSTOM_INDEX','ZSMPL_FM_GET_MAT_DETAIL','ZSMPL_FM_GET_EXPIRY_DATE',
          'ZSMPL_FM_CHECK_CHAR_DATE','ZSMPL_FM_CREATE_PROC_MSG',
          'ZSMPL_FM_GET_ASSIGNED_EQUI_EBR','/SMPL/PPPI_FM_GET_DATE_TIME',
          '/SMPL/PPPI_FM_SIG_ADD_DB_CB','/SMPL/PPPI_FM_SIG_POPULATE_CB','/SMPL/PPPI_FM_SIG_VALIDATION',
          '/SMPL/PPPI_FM_VALI_SUPE_SIG','/SMPL/PPPI_FM_INITIAL_ACTIVE','/SMPL/MBR_DEP_ADD_PERFORM',
          '/SMPL/MBR_DEP_CHECK_ACTIVE']

os.makedirs(OUTDIR, exist_ok=True)
shutil.copy(DONOR, OUT)
doc=Document(OUT); body=doc.element.body
children=list(body)
sectPr=children[-1] if children[-1].tag==qn('w:sectPr') else None
purpose=next(el for el in children if el.tag==qn('w:p') and is_heading(el) and heading_text(el).strip()=='Purpose')
anchor=copy.deepcopy(purpose)
start=children.index(purpose); end=children.index(sectPr) if sectPr is not None else len(children)
for el in children[start:end]: body.remove(el)

bid=B.max_bookmark_id(body)+1
def add(el):
    if sectPr is not None: sectPr.addprevious(el)
    else: body.append(el)
def H(t):
    global bid
    add(B.make_heading_like(anchor,t,bid)); bid+=1
def P(t): add(B.mk_plain(t))
def Bul(t): add(B.mk_bullet(t))
def BL(): add(B.mk_blank())
def TBL(tmpl, headers, rows):
    t=B.mk_table(tmpl, rows); hdr=t.findall(qn('w:tr'))[0]; tcs=hdr.findall(qn('w:tc'))
    for i,h in enumerate(headers): B.set_cell_text(tcs[i], h)
    add(t)
def T3(headers, rows): TBL(B.T_PTABLE, headers, rows)
def T2(headers, rows): TBL(B.T_ETABLE, headers, rows)
def IMG(path, width_in=6.5):
    para=doc.add_paragraph(); para.alignment=WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(path, width=Inches(width_in))
    el=para._p; el.getparent().remove(el); return el

# ---------------- content ----------------
H('Purpose')
P('The purpose of this document is to outline the design and configuration of the SMPL: VI Treatment Vessel Setup '
  'XStep. This XStep sets up the Virus Inactivation (VI) Treatment Vessel, guiding the operator through the '
  'vessel-type and filter decisions with dropdowns that conditionally activate the relevant sections: it records the '
  'bag or tank vessel, the optional vent filter and the required product filter(s), captures the vessel tare weight, '
  'confirms which filters were attached when the tare was obtained, and applies the appropriate SAP goods-issue '
  'consumption and electronic signatures.')

H('Overview')
P('The SMPL: VI Treatment Vessel Setup XStep presents the operator with a series of Yes/No selection dropdowns that '
  'drive the conditional activation of the sections beneath them, so the operator completes only the sections that '
  'apply. The operator first indicates whether the incoming Affinity product vessel will be reused as the VI '
  'Treatment Vessel - if so, the entire step is not applicable. Otherwise the Bag Required? and Tank Required? '
  'dropdowns activate the Bag Information or Tank Information section, the Vent Filter Required? dropdown activates the '
  'Vent Filter table, and the required product filters are recorded in the Product Filter table.')
P('For the bag, vent filter and product filter lines, the operator enters the Part No., which triggers an input '
  'validation that returns and displays the Material Description, and the Batch No., which triggers an input '
  'validation that returns and displays the Exp. Date; the autoclave cycle and autoclave expiration date are recorded '
  'and the expiration date is validated. When a bag or filter line is signed, a goods-issue process message (Z_PICONS, '
  'movement type 261 - order consumption) is raised to consume the material into SAP, and a per-row Performed By '
  'signature is captured. The Tank Information section records the Tank ID, CIP batch and expiration, the pressure '
  'test ID with a Get Pressure Date/Time stamp, and the SIP batch. A Get Scale / Balance button retrieves the '
  'assigned balance ID and its calibration due date; the operator records the vessel tare weight and confirms, via '
  'three Yes/No dropdowns, whether the product filter, the vent filter, or no filters (vessel is a bag) were attached '
  'when the tare was obtained. The step is closed with a Recorded By / Date signature.')
P('Each gate dropdown is a restricted Yes/No characteristic (ZSMPL_CHAR_YES_NO) captured to a local value '
  '(LV_DROP). The section it controls carries a set of setup Commands - TABLE.ACTIVATE / TABLE.DEACTIVATE on the '
  'table (and PROC_INSTR.ACTIVATE / PROC_INSTR.DEACTIVATE at the grouping level), plus LOCK / UNLOCK - whose Formula '
  'triggers evaluate the gate value: the section deactivates when the gate is No (LV_DROP = 2) and activates when it '
  'is Yes (LV_DROP <> 2), combined with the standard active-state logic (LV_ACTIVE / IV_ACT_FLP). This is the same '
  'conditional-activation pattern as the SMPL: EFS Additional Staging XStep. Because these Commands must sit in the '
  'same XStep as the section they act on, the whole vessel set-up is built as one XStep rather than a stack of '
  'separate reusable blocks. /SMPL/PPPI_FM_INITIAL_ACTIVE governs the step’s base initial-active state. All function '
  'modules used by this XStep already exist in DE1 100; no new development is required.')

H('Reasons for developing')
P('The SMPL: VI Treatment Vessel Setup XStep was developed to replace the manual "Transfer Vessel Information and '
  'Set-Up" record in the Virus Inactivation paper batch record. Driving the vessel-type, vent-filter and '
  'filters-attached branching through dropdowns that conditionally activate the relevant sections enforces the paper '
  'record’s "if bag / if tank / if applicable" logic electronically and prevents recording inapplicable data; '
  'the Part No. and Batch No. input validations remove transcription error and enforce valid part/batch '
  'combinations; raising the goods-issue process message guarantees SAP consumption of the bag and filters without a '
  'separate manual transaction; retrieving the balance as assigned equipment enforces calibration; and the '
  'electronic Performed-By and Recorded-By signatures enforce the data-integrity and verification requirements of the '
  'process. A single reusable XStep replaces a multi-part vessel-setup form.')

H('Authorization')
P('Access to the XStep and the ability to enter data and apply signatures is controlled by the standard SiMPL EBR '
  'security model. The operator must hold the PFCG role assigned to the relevant process order / control recipe. '
  'Performed-By and Recorded-By signatures are captured using the SAP digital signature method configured for the EBR.')

H('Assumptions/ Dependencies')
Bul('The generic SiMPL indexing, material/expiry/date validation, equipment, date-time, goods-issue, signature and '
    'activation function modules (ZSMPL_FM_CUSTOM_INDEX, ZSMPL_FM_GET_MAT_DETAIL, ZSMPL_FM_GET_EXPIRY_DATE, '
    'ZSMPL_FM_CHECK_CHAR_DATE, ZSMPL_FM_CREATE_PROC_MSG, ZSMPL_FM_GET_ASSIGNED_EQUI_EBR, '
    '/SMPL/PPPI_FM_GET_DATE_TIME, /SMPL/PPPI_FM_SIG_ADD_DB_CB, '
    '/SMPL/PPPI_FM_SIG_POPULATE_CB, /SMPL/PPPI_FM_SIG_VALIDATION, /SMPL/PPPI_FM_VALI_SUPE_SIG, '
    '/SMPL/PPPI_FM_INITIAL_ACTIVE, /SMPL/MBR_DEP_ADD_PERFORM, /SMPL/MBR_DEP_CHECK_ACTIVE) exist and are active in '
    'the target system.')
Bul('The CT04 characteristic ZSMPL_CHAR_YES_NO = { Yes, No } exists (build deliverable if not) and backs every '
    'Yes/No dropdown - the Use-Affinity-Vessel?, Bag Required?, Tank Required? and Vent Filter Required? gates, the '
    'Mixer Drive / Top Base / Magnetic Clamp attachment dropdowns, and the Product Filter / Vent Filter / '
    'No-filters(bag) filters-attached dropdowns.')
Bul('Conditional activation is configured with setup Commands, not a function module: each gated section carries '
    'TABLE.ACTIVATE / TABLE.DEACTIVATE (and PROC_INSTR.ACTIVATE / PROC_INSTR.DEACTIVATE, LOCK / UNLOCK) whose Formula '
    'triggers evaluate the gate value LV_DROP - deactivate when LV_DROP = 2 (No), activate when LV_DROP <> 2 (Yes), '
    'combined with LV_ACTIVE / IV_ACT_FLP (the SMPL: EFS Additional Staging pattern). /SMPL/PPPI_FM_INITIAL_ACTIVE '
    'sets the step’s base initial-active state.')
Bul('The parts entered exist in the material master (MAKT/MARA) and the batches exist in the batch master (MCH1) '
    'with a shelf-life expiration date; the bag and filters are acceptable per Section 4 (bolded part numbers '
    'require autoclaving).')
Bul('The plant, storage location and movement type (261) are configured so the Z_PICONS goods-issue process message '
    'posts the order consumption for the bag and filters.')
Bul('The scale / balance is maintained as assigned equipment (EQTYP B) with a calibration due date; the Autoclave '
    'Exp. Date is recorded in yyyy.mm.dd format, or "N/A" where an autoclave expiry does not apply.')
Bul('The Treatment Vessel is labelled per SOP-0107056; a tank’s CIP expiration is carried from the cleaning '
    'record and the pressure test / SIP records exist for the tank.')

H('Validation Checks')
P('The following validations are applied within the XStep:')
T3(['Field','Validation','Function Module'],
   [['Use Affinity Vessel? / Bag Required? / Tank Required? / Vent Filter Required?',
     'Yes/No dropdown (ZSMPL_CHAR_YES_NO -> LV_DROP); the value fires the section’s TABLE / PROC_INSTR ACTIVATE / DEACTIVATE commands',
     'Setup Commands (Formula trigger on LV_DROP) - not an FM'],
    ['Part No. (bag / vent filter / product filter)','Must be a valid SAP material; returns the Material Description','ZSMPL_FM_GET_MAT_DETAIL'],
    ['Material Description','Read-only output populated from the Part No. validation','ZSMPL_FM_GET_MAT_DETAIL'],
    ['Batch No. (bag / vent filter / product filter)','Must be a valid material/batch combination; returns the Exp. Date; not expired','ZSMPL_FM_GET_EXPIRY_DATE'],
    ['Exp. Date','Read-only output populated from the Batch No. validation','ZSMPL_FM_GET_EXPIRY_DATE'],
    ['Autoclave Exp. Date','Must be a valid date in yyyy.mm.dd format (or "N/A") and not in the past','ZSMPL_FM_CHECK_CHAR_DATE'],
    ['Get Scale / Balance','Retrieves and validates the assigned balance ID and its calibration due date','ZSMPL_FM_GET_ASSIGNED_EQUI_EBR'],
    ['Get Pressure Date/Time','Stamps the Pressure Test Date and Time','/SMPL/PPPI_FM_GET_DATE_TIME'],
    ['SAP consumption (bag / filters)','Goods-issue process message (Z_PICONS, movement type 261) raised when a line is signed','ZSMPL_FM_CREATE_PROC_MSG'],
    ['Performed By','Mandatory electronic signature per bag / filter row','/SMPL/PPPI_FM_SIG_ADD_DB_CB'],
    ['Recorded By / Date','Mandatory recording (verification) signature at step level','/SMPL/PPPI_FM_VALI_SUPE_SIG']])
BL()

H('XStep Layout Design')
P('The XStep is rendered with an instruction panel followed by gated sections. The instruction reads: "Set up the VI '
  'Treatment Vessel. The selection dropdowns below drive conditional activation - each gate activates or deactivates '
  'the section beneath it. First decide whether the incoming Affinity product vessel is reused as the VI Treatment '
  'Vessel (if so, all of Section 8 is N/A). Otherwise obtain and label the vessel (AZD0543 / VI Treatment Vessel / '
  'Batch / Part / Initials / Date), complete the Bag or Tank section, record the vent and product filters, and '
  'capture the tare weight with the filters-attached confirmation. Filters post via Goods Issue (Z_PICONS, movement '
  'type 261)." The step captures the following:')
Bul('Use Affinity Product Vessel as VI Treatment Vessel? - Yes/No dropdown (-> LV_DROP); its value fires the '
    'step’s DEACTIVATE commands so the whole step is N/A when the Affinity vessel is reused.')
Bul('Bag Required? - Yes/No dropdown; its value fires the Bag Information table’s TABLE.ACTIVATE / DEACTIVATE commands.')
Bul('Bag Information (goods issue) - Bag Part No. (entry -> Material Description), Batch No. (entry -> Exp. Date), '
    'Serial No. (entry), Performed By (signature per row); signing raises the Z_PICONS consumption.')
Bul('Mixer Drive attached? / Top Base attached? / Magnetic Clamp attached? - three Yes/No dropdowns.')
Bul('Tank Required? - Yes/No dropdown; its value fires the Tank Information table’s TABLE.ACTIVATE / DEACTIVATE commands.')
Bul('Tank Information - Tank ID (entry), CIP Batch ID (entry), CIP Exp. Date/Time (read-only, carried from the '
    'cleaning record), Pressure Test ID (entry), Get Pressure Date/Time (button) -> Pressure Test Date and Pressure '
    'Test Time (read-only, stamped), SIP Batch ID (entry).')
Bul('Vent Filter Required? - Yes/No dropdown; its value fires the Vent Filter table’s TABLE.ACTIVATE / DEACTIVATE commands.')
Bul('Vent Filter Information (goods issue) - Filter Position, Part No. (entry -> Material Description), Batch No. '
    '(entry -> Exp. Date), Serial No. (entry), Autoclave Cycle No. (entry), Autoclave Exp. (entry, date-validated), '
    'Performed By (signature per row).')
Bul('Product Filter Information (goods issue, Add Row, # index) - same columns as the Vent Filter table with an '
    'automatic # row index; one or more product filters, each with a Performed By signature.')
Bul('Get Scale / Balance (button) -> Scale / Balance ID and Calibration Due Date (read-only, retrieved).')
Bul('Tare Weight of VI Treatment Vessel (kg = L) - entry.')
Bul('Product Filter attached? / Vent Filter attached? / No filters - vessel is a bag? - three Yes/No dropdowns '
    'confirming which filters were attached when the tare was obtained.')
P('Rows can be added to the Product Filter table with Add Row; the rows are numbered automatically. A step-level '
  'Recorded By / Date signature is captured in the footer.')
add(IMG(IMG_SRC))
BL()

H('Function Module(s)')
P('The following function modules are used by this XStep. All of them already exist in DE1 100 and are reused as-is; '
  'no new function module is required.')
for fm in FM_ORDER:
    for el in B.lib_fm_elems(fm): add(el)

H('Pseudocode')
for fm in FM_ORDER:
    for el in B.lib_pseudo_elems(fm): add(el)

H('Configuration Specifications')
P('The XStep is configured using the following parameters:')
T3(['Parameter','Value / Configuration','Description'],
   [['Header','VI Treatment Vessel Setup','Header text displayed at the top of the step (configurable in SiMPL MBR).'],
    ['Instruction','(instruction text)','The instruction text displayed to the operator (configurable in SiMPL MBR).'],
    ['Conditional activation','Setup Commands + Formula trigger','Each gated section carries TABLE.ACTIVATE / TABLE.DEACTIVATE (and PROC_INSTR.ACTIVATE / DEACTIVATE, LOCK / UNLOCK); the ACTIVATE/DEACTIVATE Formula triggers evaluate the gate value LV_DROP - e.g. DEACTIVATE fires on ( LV_ACTIVE <> 1 AND IV_ACT_FLP <> 1 ) OR ( LV_ACTIVE = 1 AND IV_ACT_FLP = 1 ) OR ( LV_DROP = 2 ). Not a function module.'],
    ['Use Affinity Vessel?','Dropdown (CT04) -> LV_DROP','ZSMPL_CHAR_YES_NO = Yes / No; No/other value fires the step’s DEACTIVATE commands (whole step N/A).'],
    ['Bag Required?','Dropdown (CT04) -> LV_DROP','ZSMPL_CHAR_YES_NO = Yes / No; drives the Bag Information TABLE.ACTIVATE / DEACTIVATE commands.'],
    ['Bag Part No.','Entry, input validation','Validated by ZSMPL_FM_GET_MAT_DETAIL, which returns the Material Description.'],
    ['Bag Batch No.','Entry, input validation','Validated by ZSMPL_FM_GET_EXPIRY_DATE, which returns the Exp. Date.'],
    ['Mixer Drive / Top Base / Magnetic Clamp attached?','Dropdown (CT04) x3','ZSMPL_CHAR_YES_NO = Yes / No (mixer-bag attachments).'],
    ['Tank Required?','Dropdown (CT04) -> LV_DROP','ZSMPL_CHAR_YES_NO = Yes / No; drives the Tank Information TABLE.ACTIVATE / DEACTIVATE commands.'],
    ['Tank ID / CIP Batch ID / Pressure Test ID / SIP Batch ID','Entry fields','Free-text entry for the tank and cleaning/sterilization references.'],
    ['CIP Exp. Date/Time','Read-only output','Carried from the CIP cleaning record.'],
    ['Get Pressure Date/Time','Button','Stamps the Pressure Test Date and Time (/SMPL/PPPI_FM_GET_DATE_TIME).'],
    ['Vent Filter Required?','Dropdown (CT04) -> LV_DROP','ZSMPL_CHAR_YES_NO = Yes / No; drives the Vent Filter TABLE.ACTIVATE / DEACTIVATE commands.'],
    ['Filter Part No. / Batch No.','Entry, input validation','ZSMPL_FM_GET_MAT_DETAIL (description) and ZSMPL_FM_GET_EXPIRY_DATE (expiry).'],
    ['Autoclave Exp. Date','Entry, date validator','Validated by ZSMPL_FM_CHECK_CHAR_DATE (yyyy.mm.dd or "N/A", not in past).'],
    ['Get Scale / Balance','Button (equipment)','Retrieves + validates the assigned balance ID + calibration due (ZSMPL_FM_GET_ASSIGNED_EQUI_EBR, EQTYP B).'],
    ['Tare Weight','Entry field','Numeric entry for the vessel tare weight (kg = L).'],
    ['Filters attached (Product / Vent / None-bag)','Dropdown (CT04) x3','ZSMPL_CHAR_YES_NO = Yes / No; which filters were attached at tare.'],
    ['Goods Issue','Z_PICONS process message','Raised by ZSMPL_FM_CREATE_PROC_MSG on signing a bag/filter line (movement type 261).'],
    ['Performed By','Signature column','Per-row Performed-By signature on the bag/filter tables (/SMPL/PPPI_FM_SIG_ADD_DB_CB).'],
    ['Recorded By / Date','Step-level signature','Recording/verification signature (/SMPL/PPPI_FM_VALI_SUPE_SIG).'],
    ['Add Row','Enabled (Product Filter table)','Rows numbered automatically by ZSMPL_FM_CUSTOM_INDEX.']])
BL()

H('Test Scenarios')
T3(['#','Scenario','Expected Result'],
   [['1','Set Use Affinity Vessel? = Yes.','The step’s DEACTIVATE commands fire (via LV_DROP) so the whole step is N/A and the operator continues to the next section.'],
    ['2','Set Bag Required? = Yes / No.','Yes (LV_DROP <> 2) fires TABLE.ACTIVATE so the Bag Information section activates for entry; No (LV_DROP = 2) fires TABLE.DEACTIVATE so it stays deactivated.'],
    ['3','Set Tank Required? = Yes.','The Tank Information section activates for entry; the Bag section can remain deactivated.'],
    ['4','Enter a valid Part No. for a bag or filter.','The Material Description is returned and displayed read-only (ZSMPL_FM_GET_MAT_DETAIL).'],
    ['5','Enter a valid Batch No. for the entered material.','The Exp. Date is returned and displayed read-only (ZSMPL_FM_GET_EXPIRY_DATE).'],
    ['6','Enter a Batch No. whose expiration date is in the past.','The entry is rejected with EXPIRY_DATE_IN_PAST.'],
    ['7','Enter an Autoclave Exp. Date with an invalid month (e.g. 2027.13.01).','The entry is rejected with "Incorrect date format entered - use yyyy.mm.dd" (INCORRECT_DATE_FORMAT).'],
    ['8','Enter "N/A" for the Autoclave Exp. Date.','The entry is accepted (no autoclave expiry applies).'],
    ['9','Press Get Pressure Date/Time in the Tank section.','The Pressure Test Date and Time are stamped read-only (/SMPL/PPPI_FM_GET_DATE_TIME).'],
    ['10','Press Get Scale / Balance.','The assigned balance ID and its calibration due date are retrieved and displayed (ZSMPL_FM_GET_ASSIGNED_EQUI_EBR).'],
    ['11','Apply the Performed-By signature on a completed bag/filter row.','The row is signed and a Z_PICONS goods-issue process message (movement type 261) is raised to consume the material into SAP (ZSMPL_FM_CREATE_PROC_MSG).'],
    ['12','Add a second row in the Product Filter table with Add Row.','A new numbered row is added and can be completed independently (ZSMPL_FM_CUSTOM_INDEX).'],
    ['13','Attempt to complete the step without the Recorded By / Date signature.','The step cannot be completed; the recording signature is mandatory.'],
    ['14','Re-open the completed step.','The recorded values, validation outputs, stamps and signatures are displayed and retained.']])
BL()

H('Document References')
T2(['Reference No.','Document Title'],
   [['1','AZD0543 - Virus Inactivation (Low pH) Manufacturing Directions (PN 8012441)'],
    ['2','AZD0543 - Concentration & Diafiltration Manufacturing Directions (PN 8012475)'],
    ['3','SiMPL XStep Library - Process Manufacturing (DE1 100)'],
    ['4','SMPL: Room/Equipment Assign XStep Design Specification V1.3 (reference template)']])
BL()

H('Revision History')
T3(['Version No.','Description','Date'],
   [['1.0','Initial document','2026-07-23']])

# ---------------- fixes + save ----------------
B.neutralise_empty_headings(body)
B.set_update_fields(doc)
doc.save(OUT)

# title-canvas + header/footer text: rename the XStep
tmp=OUT+'.tmp'
with zipfile.ZipFile(OUT) as z: data={n:z.read(n) for n in z.namelist()}
for n in list(data):
    if re.match(r'word/(document|header\d+|footer\d+)\.xml$', n):
        x=data[n].decode('utf-8')
        x2=x.replace('Three Variable Calc','VI Treatment Vessel Setup')
        if x2!=x: data[n]=x2.encode('utf-8')
with zipfile.ZipFile(tmp,'w',zipfile.ZIP_DEFLATED) as z:
    for n,b in data.items(): z.writestr(n,b)
shutil.move(tmp,OUT)
sys.stdout.flush()
print('built',OUT)
